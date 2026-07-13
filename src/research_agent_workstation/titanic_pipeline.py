from __future__ import annotations

import argparse
import json
import re
import time
from datetime import datetime
from pathlib import Path
from typing import Any

import pandas as pd
import yaml
from sklearn.compose import ColumnTransformer
from sklearn.ensemble import ExtraTreesClassifier, GradientBoostingClassifier, RandomForestClassifier
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import accuracy_score, f1_score
from sklearn.model_selection import StratifiedKFold, cross_val_score, train_test_split
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import OneHotEncoder, StandardScaler

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
except ImportError:  # DOCX export is optional; Markdown/JSON remain the source of truth.
    Document = None
    qn = None
    Pt = None
    RGBColor = None


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Run the local Titanic Kaggle workflow.")
    parser.add_argument("--config", default="configs/titanic.yaml")
    parser.add_argument("--output-dir", default="experiments/titanic")
    parser.add_argument("--random-state", type=int, default=42)
    return parser.parse_args()


def load_config(path: Path) -> dict[str, Any]:
    return yaml.safe_load(path.read_text(encoding="utf-8"))


def normalize_title(title: str) -> str:
    title = title.strip()
    mapping = {
        "Mlle": "Miss",
        "Ms": "Miss",
        "Mme": "Mrs",
        "Lady": "Rare",
        "Countess": "Rare",
        "Capt": "Rare",
        "Col": "Rare",
        "Don": "Rare",
        "Dr": "Rare",
        "Major": "Rare",
        "Rev": "Rare",
        "Sir": "Rare",
        "Jonkheer": "Rare",
        "Dona": "Rare",
    }
    return mapping.get(title, title)


def engineer_features(df: pd.DataFrame) -> pd.DataFrame:
    result = df.copy()

    result["Title"] = (
        result["Name"]
        .fillna("")
        .map(lambda value: re.search(r",\s*([^\.]+)\.", value).group(1) if re.search(r",\s*([^\.]+)\.", value) else "Unknown")
        .map(normalize_title)
    )
    result["FamilySize"] = result["SibSp"].fillna(0) + result["Parch"].fillna(0) + 1
    result["IsAlone"] = (result["FamilySize"] == 1).astype(int)
    result["Deck"] = result["Cabin"].fillna("U").astype(str).str[0]
    result["FarePerPerson"] = result["Fare"] / result["FamilySize"].replace(0, 1)
    result["AgeMissing"] = result["Age"].isna().astype(int)
    result["FareMissing"] = result["Fare"].isna().astype(int)

    return result.drop(columns=["PassengerId", "Name", "Ticket", "Cabin"], errors="ignore")


def build_preprocessor(x: pd.DataFrame) -> ColumnTransformer:
    numeric_cols = x.select_dtypes(include="number").columns.tolist()
    categorical_cols = [col for col in x.columns if col not in numeric_cols]

    numeric_pipe = Pipeline(
        steps=[
            ("imputer", SimpleImputer(strategy="median")),
            ("scaler", StandardScaler()),
        ]
    )
    categorical_pipe = Pipeline(
        steps=[
            ("imputer", SimpleImputer(strategy="most_frequent")),
            ("encoder", OneHotEncoder(handle_unknown="ignore")),
        ]
    )

    return ColumnTransformer(
        transformers=[
            ("num", numeric_pipe, numeric_cols),
            ("cat", categorical_pipe, categorical_cols),
        ]
    )


def data_quality_report(train: pd.DataFrame, test: pd.DataFrame, sample: pd.DataFrame, target: str) -> dict[str, Any]:
    train_features = [col for col in train.columns if col != target]
    test_features = test.columns.tolist()
    missing_train = train.isna().mean().sort_values(ascending=False).head(10)
    missing_test = test.isna().mean().sort_values(ascending=False).head(10)
    target_counts = train[target].value_counts(normalize=True).round(4).to_dict()

    return {
        "train_rows": int(len(train)),
        "train_columns": int(train.shape[1]),
        "test_rows": int(len(test)),
        "test_columns": int(test.shape[1]),
        "sample_submission_rows": int(len(sample)),
        "sample_submission_columns": sample.columns.tolist(),
        "target": target,
        "target_distribution": {str(k): float(v) for k, v in target_counts.items()},
        "train_test_feature_columns_match": train_features == test_features,
        "missing_train_top10": {k: round(float(v), 4) for k, v in missing_train.items()},
        "missing_test_top10": {k: round(float(v), 4) for k, v in missing_test.items()},
    }


def build_task_scaffold(config: dict[str, Any], quality: dict[str, Any]) -> dict[str, Any]:
    data_cfg = config["data"]
    scaffold_cfg = config.get("scaffold", {})
    return {
        "task": config["task"],
        "inputs": {
            "overview": data_cfg["overview"],
            "train": data_cfg["train"],
            "test": data_cfg["test"],
            "sample_submission": data_cfg["sample_submission"],
        },
        "data_snapshot": {
            "train_rows": quality["train_rows"],
            "test_rows": quality["test_rows"],
            "sample_submission_rows": quality["sample_submission_rows"],
            "target_distribution": quality["target_distribution"],
            "missing_train_top10": quality["missing_train_top10"],
            "missing_test_top10": quality["missing_test_top10"],
        },
        "validation_strategy": scaffold_cfg.get("validation_strategy", "stratified validation"),
        "time_budget_minutes": scaffold_cfg.get("time_budget_minutes", 10),
        "candidate_models": scaffold_cfg.get("first_stage_models", list(candidate_models(42).keys())),
        "feature_plan": [
            "Extract passenger title from Name.",
            "Create FamilySize and IsAlone.",
            "Create Deck from Cabin with U as missing marker.",
            "Create FarePerPerson and missing indicators.",
            "Use median imputation for numeric features and most-frequent imputation for categorical features.",
        ],
        "risk_points": scaffold_cfg.get("risk_points", []),
        "stage_plan": config["workflow"],
    }


def write_scaffold(output_dir: Path, scaffold: dict[str, Any]) -> None:
    (output_dir / "task_scaffold.json").write_text(
        json.dumps(scaffold, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    lines = [
        "# Titanic 任务脚手架",
        "",
        "## 任务",
        "",
        f"- 比赛：{scaffold['task']['competition']}",
        f"- 类型：{scaffold['task']['type']}",
        f"- 目标列：{scaffold['task']['target']}",
        f"- 指标：{scaffold['task']['metric']}",
        "",
        "## 输入文件",
        "",
        *[f"- {name}: `{path}`" for name, path in scaffold["inputs"].items()],
        "",
        "## 验证方案",
        "",
        f"- {scaffold['validation_strategy']}",
        f"- 时间预算：{scaffold['time_budget_minutes']} 分钟",
        "",
        "## 候选模型",
        "",
        *[f"- {name}" for name in scaffold["candidate_models"]],
        "",
        "## 特征计划",
        "",
        *[f"- {item}" for item in scaffold["feature_plan"]],
        "",
        "## 风险点",
        "",
        *[f"- {item}" for item in scaffold["risk_points"]],
    ]
    (output_dir / "task_scaffold.md").write_text("\n".join(lines), encoding="utf-8-sig")


def build_stage_audit(
    config: dict[str, Any],
    quality: dict[str, Any],
    evaluation: dict[str, Any],
    submission: dict[str, Any],
    accepted: bool,
    output_dir: Path,
) -> dict[str, Any]:
    best_model = evaluation["best_model"]
    best_metrics = evaluation["model_results"][best_model]
    thresholds = config["thresholds"]

    stages = [
        {
            "stage": "task_understanding",
            "owner_role": "Reader/Planner",
            "status": "passed",
            "evidence": [
                config["data"]["overview"],
                str(output_dir / "task_scaffold.json"),
            ],
            "checks": {
                "task_name": config["task"]["name"],
                "target": config["task"]["target"],
                "metric": config["task"]["metric"],
                "input_files_declared": True,
            },
        },
        {
            "stage": "preliminary_eda",
            "owner_role": "Analyst",
            "status": "passed",
            "evidence": [str(output_dir / "data_quality.json")],
            "checks": {
                "train_rows": quality["train_rows"],
                "test_rows": quality["test_rows"],
                "target_distribution_recorded": bool(quality["target_distribution"]),
                "missing_values_recorded": bool(quality["missing_train_top10"]) and bool(quality["missing_test_top10"]),
            },
        },
        {
            "stage": "data_quality_check",
            "owner_role": "Reviewer",
            "status": "passed" if quality["train_test_feature_columns_match"] else "failed",
            "evidence": [str(output_dir / "data_quality.json")],
            "checks": {
                "train_test_feature_columns_match": quality["train_test_feature_columns_match"],
                "sample_submission_rows": quality["sample_submission_rows"],
                "sample_submission_columns": quality["sample_submission_columns"],
            },
        },
        {
            "stage": "feature_engineering",
            "owner_role": "Developer",
            "status": "passed",
            "evidence": [str(output_dir / "task_scaffold.json")],
            "checks": {
                "engineered_feature_groups": 5,
                "uses_title_family_deck_missing_features": True,
                "uses_imputation_and_encoding_pipeline": True,
            },
        },
        {
            "stage": "model_validation",
            "owner_role": "Developer/Reviewer",
            "status": "passed" if best_metrics["cv_accuracy_mean"] >= thresholds["min_validation_accuracy"] else "failed",
            "evidence": [str(output_dir / "model_results.json")],
            "checks": {
                "best_model": best_model,
                "cv_accuracy_mean": best_metrics["cv_accuracy_mean"],
                "holdout_accuracy": best_metrics["holdout_accuracy"],
                "min_validation_accuracy": thresholds["min_validation_accuracy"],
                "candidate_model_count": len(evaluation["model_results"]),
            },
        },
        {
            "stage": "submission_generation",
            "owner_role": "Developer/Reviewer",
            "status": "passed" if submission["valid"] else "failed",
            "evidence": [submission["path"]],
            "checks": {
                "rows_match": submission["rows_match"],
                "columns_match": submission["columns_match"],
                "missing_predictions": submission["missing_predictions"],
                "allowed_values_only": submission["allowed_values_only"],
                "prediction_distribution": submission["prediction_distribution"],
            },
        },
        {
            "stage": "report_and_review",
            "owner_role": "Summarizer/Reviewer",
            "status": "passed" if accepted else "failed",
            "evidence": [
                str(output_dir / "titanic_local_report.md"),
                str(output_dir / "titanic_local_report.docx"),
            ],
            "checks": {
                "local_gate_expected": accepted,
                "validation_gate_written_by": "scripts/validate_titanic_experiment.py",
                "markdown_report_path": str(output_dir / "titanic_local_report.md"),
                "docx_report_path": str(output_dir / "titanic_local_report.docx"),
            },
        },
    ]

    failed = [stage["stage"] for stage in stages if stage["status"] != "passed"]
    return {
        "task": config["task"],
        "experiment_dir": str(output_dir),
        "workflow_version": "v2_stage_audit",
        "local_only": True,
        "kaggle_api_submission": False,
        "stages": stages,
        "all_stages_passed": not failed,
        "failed_stages": failed,
        "next_actions": [
            "Configure Kaggle API token before official download/submission.",
            "Run the same staged workflow on a second tabular Kaggle task to test transferability.",
            "Keep GPU/server work disabled until a task has a clear compute benefit.",
        ],
    }


def write_stage_audit(output_dir: Path, audit: dict[str, Any]) -> None:
    (output_dir / "workflow_stage_audit.json").write_text(
        json.dumps(audit, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    lines = [
        "# Titanic 阶段化工作流审计",
        "",
        "## 总结",
        "",
        f"- 实验目录：`{audit['experiment_dir']}`",
        f"- 工作流版本：`{audit['workflow_version']}`",
        f"- 全部阶段通过：`{audit['all_stages_passed']}`",
        f"- 官方 Kaggle 提交：`{audit['kaggle_api_submission']}`",
        "",
        "## 阶段检查",
        "",
    ]
    for stage in audit["stages"]:
        lines.extend(
            [
                f"### {stage['stage']}",
                "",
                f"- 角色：{stage['owner_role']}",
                f"- 状态：`{stage['status']}`",
                f"- 证据：{', '.join(f'`{item}`' for item in stage['evidence'])}",
                f"- 检查项：`{json.dumps(stage['checks'], ensure_ascii=False)}`",
                "",
            ]
        )

    lines.extend(["## 下一步", ""])
    lines.extend(f"- {item}" for item in audit["next_actions"])
    (output_dir / "workflow_stage_audit.md").write_text("\n".join(lines), encoding="utf-8-sig")


def candidate_models(random_state: int) -> dict[str, Any]:
    return {
        "logistic_regression": LogisticRegression(max_iter=2000, random_state=random_state),
        "random_forest": RandomForestClassifier(
            n_estimators=300,
            max_depth=8,
            min_samples_leaf=2,
            random_state=random_state,
            n_jobs=-1,
        ),
        "extra_trees": ExtraTreesClassifier(
            n_estimators=400,
            max_depth=8,
            min_samples_leaf=2,
            random_state=random_state,
            n_jobs=-1,
        ),
        "gradient_boosting": GradientBoostingClassifier(random_state=random_state),
    }


def evaluate_models(x: pd.DataFrame, y: pd.Series, random_state: int) -> tuple[dict[str, Any], Pipeline]:
    x_train, x_valid, y_train, y_valid = train_test_split(
        x,
        y,
        test_size=0.2,
        random_state=random_state,
        stratify=y,
    )
    cv = StratifiedKFold(n_splits=5, shuffle=True, random_state=random_state)
    preprocessor = build_preprocessor(x)
    results: dict[str, Any] = {}
    best_name = ""
    best_score = -1.0
    best_pipeline: Pipeline | None = None

    for name, model in candidate_models(random_state).items():
        pipeline = Pipeline(
            steps=[
                ("preprocessor", preprocessor),
                ("model", model),
            ]
        )
        start = time.time()
        cv_scores = cross_val_score(pipeline, x, y, cv=cv, scoring="accuracy", n_jobs=-1)
        pipeline.fit(x_train, y_train)
        valid_pred = pipeline.predict(x_valid)
        elapsed = round(time.time() - start, 4)
        metrics = {
            "cv_accuracy_mean": round(float(cv_scores.mean()), 6),
            "cv_accuracy_std": round(float(cv_scores.std()), 6),
            "holdout_accuracy": round(float(accuracy_score(y_valid, valid_pred)), 6),
            "holdout_macro_f1": round(float(f1_score(y_valid, valid_pred, average="macro")), 6),
            "seconds": elapsed,
        }
        results[name] = metrics
        if metrics["cv_accuracy_mean"] > best_score:
            best_score = metrics["cv_accuracy_mean"]
            best_name = name
            best_pipeline = pipeline

    if best_pipeline is None:
        raise RuntimeError("No model was trained.")

    best_pipeline.fit(x, y)
    return {"best_model": best_name, "model_results": results}, best_pipeline


def make_submission(best_pipeline: Pipeline, test_features: pd.DataFrame, sample: pd.DataFrame, output_dir: Path) -> dict[str, Any]:
    predictions = best_pipeline.predict(test_features).astype(int)
    submission = sample.copy()
    prediction_column = sample.columns[1]
    submission[prediction_column] = predictions

    path = output_dir / "submission.csv"
    submission.to_csv(path, index=False)

    checks = {
        "path": str(path),
        "rows_match": len(submission) == len(sample),
        "columns_match": submission.columns.tolist() == sample.columns.tolist(),
        "missing_predictions": int(submission[prediction_column].isna().sum()),
        "allowed_values_only": sorted(submission[prediction_column].dropna().unique().tolist()) in ([0], [1], [0, 1]),
        "prediction_distribution": {str(k): int(v) for k, v in submission[prediction_column].value_counts().to_dict().items()},
    }
    checks["valid"] = (
        checks["rows_match"]
        and checks["columns_match"]
        and checks["missing_predictions"] == 0
        and checks["allowed_values_only"]
    )
    return checks


def write_docx_report(markdown_path: Path, docx_path: Path) -> None:
    if Document is None:
        return

    document = Document()
    font_name = "宋体"
    font_name = "\u5b8b\u4f53"
    black = RGBColor(0, 0, 0)
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "List Bullet"]:
        style = document.styles[style_name]
        style.font.name = font_name
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        style.font.color.rgb = black
        if style_name == "Normal":
            style.font.size = Pt(11)

    for raw in markdown_path.read_text(encoding="utf-8-sig").splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            paragraph = document.add_paragraph(line[2:].strip(), style="Title")
        elif line.startswith("## "):
            paragraph = document.add_paragraph(line[3:].strip(), style="Heading 1")
        elif line.startswith("- "):
            paragraph = document.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            paragraph = document.add_paragraph(line)
        for run in paragraph.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
            run.font.color.rgb = black

    document.save(docx_path)


def write_markdown_report(output_dir: Path, quality: dict[str, Any], evaluation: dict[str, Any], submission: dict[str, Any], accepted: bool) -> None:
    best = evaluation["best_model"]
    best_metrics = evaluation["model_results"][best]
    lines = [
        "# Titanic Kaggle 本地测试报告",
        "",
        "## 任务理解",
        "",
        "- 任务：Titanic 生存预测。",
        "- 类型：二分类表格数据任务。",
        "- 指标：Kaggle 官方使用 accuracy，本地用 5 折交叉验证 accuracy 和 holdout accuracy 做代理评估。",
        "- 当前说明：本机暂未配置 Kaggle API 凭据，因此本轮使用公开镜像数据完成本地闭环测试；后续可替换为官方 Kaggle API 下载。",
        "",
        "## 数据质量",
        "",
        f"- train 行列数：{quality['train_rows']} x {quality['train_columns']}",
        f"- test 行列数：{quality['test_rows']} x {quality['test_columns']}",
        f"- sample_submission 行数：{quality['sample_submission_rows']}",
        f"- 训练/测试特征列一致：{quality['train_test_feature_columns_match']}",
        f"- 目标分布：{json.dumps(quality['target_distribution'], ensure_ascii=False)}",
        f"- 训练集缺失率 Top10：{json.dumps(quality['missing_train_top10'], ensure_ascii=False)}",
        f"- 测试集缺失率 Top10：{json.dumps(quality['missing_test_top10'], ensure_ascii=False)}",
        "",
        "## 模型验证",
        "",
        f"- 最佳模型：`{best}`",
        f"- 5 折 CV accuracy：{best_metrics['cv_accuracy_mean']} ± {best_metrics['cv_accuracy_std']}",
        f"- Holdout accuracy：{best_metrics['holdout_accuracy']}",
        f"- Holdout macro-F1：{best_metrics['holdout_macro_f1']}",
        "",
        "## Submission 检查",
        "",
        f"- submission 文件：`{submission['path']}`",
        f"- 行数匹配：{submission['rows_match']}",
        f"- 列名匹配：{submission['columns_match']}",
        f"- 缺失预测数：{submission['missing_predictions']}",
        f"- 预测值合法：{submission['allowed_values_only']}",
        f"- 预测分布：{json.dumps(submission['prediction_distribution'], ensure_ascii=False)}",
        "",
        "## 验收结论",
        "",
        f"- 本地 v1 验收：{'通过' if accepted else '未通过'}",
        "- 判断依据：数据质量检查、模型验证分数、submission 格式检查和实验日志均已记录。",
        "",
        "## 下一步",
        "",
        "- 配置 Kaggle API token 后，替换为官方 Kaggle API 下载与可选真实提交。",
        "- 加入 scaffold 任务脚手架，让每个比赛先生成指标、验证方案、模型路线和风险点。",
        "- 增加 LightGBM/XGBoost/CatBoost 或 GPU 搜索任务，但前提是先保持当前闭环稳定。",
    ]
    markdown_path = output_dir / "titanic_local_report.md"
    markdown_path.write_text("\n".join(lines), encoding="utf-8-sig")
    write_docx_report(markdown_path, output_dir / "titanic_local_report.docx")


def main() -> None:
    args = parse_args()
    config = load_config(Path(args.config))
    data_cfg = config["data"]
    thresholds = config["thresholds"]

    train_path = Path(data_cfg["train"])
    test_path = Path(data_cfg["test"])
    sample_path = Path(data_cfg["sample_submission"])
    for path in (train_path, test_path, sample_path):
        if not path.exists():
            raise FileNotFoundError(f"Missing data file: {path}. Run scripts/prepare_titanic_data.py first.")

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_dir = Path(args.output_dir) / timestamp
    output_dir.mkdir(parents=True, exist_ok=True)

    train = pd.read_csv(train_path)
    test = pd.read_csv(test_path)
    sample = pd.read_csv(sample_path)
    target = config["task"]["target"]

    quality = data_quality_report(train, test, sample, target)
    scaffold = build_task_scaffold(config, quality)
    write_scaffold(output_dir, scaffold)

    x = engineer_features(train.drop(columns=[target]))
    y = train[target]
    test_features = engineer_features(test)

    evaluation, best_pipeline = evaluate_models(x, y, args.random_state)
    submission = make_submission(best_pipeline, test_features, sample, output_dir)

    best_metrics = evaluation["model_results"][evaluation["best_model"]]
    accepted = bool(
        quality["train_test_feature_columns_match"]
        and submission["valid"]
        and best_metrics["cv_accuracy_mean"] >= thresholds["min_validation_accuracy"]
    )

    artifacts = {
        "task": config["task"],
        "data_quality": quality,
        "evaluation": evaluation,
        "submission_check": submission,
        "task_scaffold": scaffold,
        "accepted": accepted,
        "thresholds": thresholds,
        "outputs": {
            "submission": submission["path"],
            "scaffold": str(output_dir / "task_scaffold.json"),
            "scaffold_markdown": str(output_dir / "task_scaffold.md"),
            "report": str(output_dir / "titanic_local_report.md"),
            "report_docx": str(output_dir / "titanic_local_report.docx"),
        },
    }
    (output_dir / "experiment_log.json").write_text(json.dumps(artifacts, ensure_ascii=False, indent=2), encoding="utf-8")
    (output_dir / "data_quality.json").write_text(json.dumps(quality, ensure_ascii=False, indent=2), encoding="utf-8")
    (output_dir / "model_results.json").write_text(json.dumps(evaluation, ensure_ascii=False, indent=2), encoding="utf-8")
    write_markdown_report(output_dir, quality, evaluation, submission, accepted)
    stage_audit = build_stage_audit(config, quality, evaluation, submission, accepted, output_dir)
    write_stage_audit(output_dir, stage_audit)

    print(json.dumps({"output_dir": str(output_dir), "accepted": accepted, "best": evaluation["best_model"], "metrics": best_metrics}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
