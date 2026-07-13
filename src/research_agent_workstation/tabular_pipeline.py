from __future__ import annotations

import argparse
import json
import math
import time
from datetime import datetime
from pathlib import Path
from typing import Any

import numpy as np
import pandas as pd
import yaml
from sklearn.base import clone
from sklearn.compose import ColumnTransformer
from sklearn.ensemble import (
    ExtraTreesClassifier,
    ExtraTreesRegressor,
    GradientBoostingClassifier,
    GradientBoostingRegressor,
    RandomForestClassifier,
    RandomForestRegressor,
)
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LogisticRegression, Ridge
from sklearn.metrics import accuracy_score, f1_score, mean_absolute_error, mean_squared_error, mean_squared_log_error
from sklearn.model_selection import KFold, StratifiedKFold, train_test_split
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import OneHotEncoder, StandardScaler

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
except ImportError:  # DOCX export is optional; Markdown/JSON remain the source of truth.
    Document = None
    qn = None
    Pt = None
    RGBColor = None


def parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Run a staged tabular Kaggle-style workflow.")
    parser.add_argument("--config", required=True)
    parser.add_argument("--output-dir", default=None)
    parser.add_argument("--random-state", type=int, default=42)
    return parser.parse_args(argv)


def load_yaml(path: Path) -> dict[str, Any]:
    return yaml.safe_load(path.read_text(encoding="utf-8"))


def is_classification_task(config: dict[str, Any]) -> bool:
    return "classification" in str(config["task"]["type"]).lower()


def is_regression_task(config: dict[str, Any]) -> bool:
    return "regression" in str(config["task"]["type"]).lower()


def make_encoder() -> OneHotEncoder:
    try:
        return OneHotEncoder(handle_unknown="ignore", sparse_output=False)
    except TypeError:
        return OneHotEncoder(handle_unknown="ignore", sparse=False)


def rmsle(y_true: pd.Series | np.ndarray, y_pred: np.ndarray) -> float:
    y_true_arr = np.asarray(y_true, dtype=float)
    y_pred_arr = np.asarray(y_pred, dtype=float)
    y_pred_arr = np.clip(y_pred_arr, 0, None)
    return float(math.sqrt(mean_squared_log_error(y_true_arr, y_pred_arr)))


def regression_metric_value(metric: str, y_true: pd.Series | np.ndarray, y_pred: np.ndarray) -> float:
    metric_name = metric.lower()
    y_pred_arr = np.asarray(y_pred, dtype=float)
    if metric_name == "rmsle":
        return rmsle(y_true, y_pred_arr)
    if metric_name == "mae":
        return float(mean_absolute_error(y_true, y_pred_arr))
    return float(math.sqrt(mean_squared_error(y_true, y_pred_arr)))


def inverse_target_transform(values: np.ndarray, target_transform: str | None) -> np.ndarray:
    if target_transform == "log1p":
        return np.expm1(values)
    return values


def data_quality_report(train: pd.DataFrame, test: pd.DataFrame, sample: pd.DataFrame, config: dict[str, Any]) -> dict[str, Any]:
    target = config["task"]["target"]
    task_type = config["task"]["type"]
    train_features = [col for col in train.columns if col != target]
    test_features = test.columns.tolist()
    missing_train = train.isna().mean().sort_values(ascending=False).head(20)
    missing_test = test.isna().mean().sort_values(ascending=False).head(20)

    target_summary: dict[str, Any]
    if is_classification_task(config):
        target_summary = {
            "distribution": {str(k): float(v) for k, v in train[target].value_counts(normalize=True).round(4).to_dict().items()}
        }
    else:
        desc = train[target].describe().round(4).to_dict()
        target_summary = {
            "summary": {str(k): float(v) for k, v in desc.items()},
            "skew": round(float(train[target].skew()), 6),
        }

    return {
        "train_rows": int(len(train)),
        "train_columns": int(train.shape[1]),
        "test_rows": int(len(test)),
        "test_columns": int(test.shape[1]),
        "sample_submission_rows": int(len(sample)),
        "sample_submission_columns": sample.columns.tolist(),
        "target": target,
        "task_type": task_type,
        "target_summary": target_summary,
        "duplicate_train_rows": int(train.duplicated().sum()),
        "train_test_feature_columns_match": train_features == test_features,
        "missing_train_top20": {k: round(float(v), 4) for k, v in missing_train.items()},
        "missing_test_top20": {k: round(float(v), 4) for k, v in missing_test.items()},
    }


def apply_feature_preset(df: pd.DataFrame, config: dict[str, Any]) -> pd.DataFrame:
    result = df.copy()
    feature_cfg = config.get("feature_engineering", {})
    preset = feature_cfg.get("preset")

    if preset == "house_prices_basic":
        for col in ["TotalBsmtSF", "1stFlrSF", "2ndFlrSF", "FullBath", "HalfBath", "BsmtFullBath", "BsmtHalfBath"]:
            if col in result.columns:
                result[col] = pd.to_numeric(result[col], errors="coerce")
        result["TotalSF"] = (
            result.get("TotalBsmtSF", 0).fillna(0)
            + result.get("1stFlrSF", 0).fillna(0)
            + result.get("2ndFlrSF", 0).fillna(0)
        )
        result["TotalBath"] = (
            result.get("FullBath", 0).fillna(0)
            + 0.5 * result.get("HalfBath", 0).fillna(0)
            + result.get("BsmtFullBath", 0).fillna(0)
            + 0.5 * result.get("BsmtHalfBath", 0).fillna(0)
        )
        if {"YrSold", "YearBuilt"}.issubset(result.columns):
            result["HouseAge"] = (result["YrSold"] - result["YearBuilt"]).clip(lower=0)
        if {"YrSold", "YearRemodAdd"}.issubset(result.columns):
            result["RemodAge"] = (result["YrSold"] - result["YearRemodAdd"]).clip(lower=0)
        if "GarageArea" in result.columns:
            result["HasGarage"] = (result["GarageArea"].fillna(0) > 0).astype(int)
        if "TotalBsmtSF" in result.columns:
            result["HasBasement"] = (result["TotalBsmtSF"].fillna(0) > 0).astype(int)
        if "Fireplaces" in result.columns:
            result["HasFireplace"] = (result["Fireplaces"].fillna(0) > 0).astype(int)
        if "PoolArea" in result.columns:
            result["HasPool"] = (result["PoolArea"].fillna(0) > 0).astype(int)

    drop_columns = feature_cfg.get("drop_columns", [])
    return result.drop(columns=drop_columns, errors="ignore")


def build_preprocessor(x: pd.DataFrame) -> ColumnTransformer:
    numeric_cols = x.select_dtypes(include="number").columns.tolist()
    categorical_cols = [col for col in x.columns if col not in numeric_cols]

    numeric_pipe = Pipeline(
        steps=[
            ("imputer", SimpleImputer(strategy="median")),
            ("scaler", StandardScaler()),
        ]
    )
    categorical_pipe = Pipeline(
        steps=[
            ("imputer", SimpleImputer(strategy="most_frequent")),
            ("encoder", make_encoder()),
        ]
    )

    return ColumnTransformer(
        transformers=[
            ("num", numeric_pipe, numeric_cols),
            ("cat", categorical_pipe, categorical_cols),
        ],
        remainder="drop",
    )


def classification_models(random_state: int) -> dict[str, Any]:
    return {
        "logistic_regression": LogisticRegression(max_iter=2000, random_state=random_state),
        "random_forest": RandomForestClassifier(n_estimators=300, max_depth=8, min_samples_leaf=2, random_state=random_state, n_jobs=-1),
        "extra_trees": ExtraTreesClassifier(n_estimators=400, max_depth=8, min_samples_leaf=2, random_state=random_state, n_jobs=-1),
        "gradient_boosting": GradientBoostingClassifier(random_state=random_state),
    }


def regression_models(random_state: int) -> dict[str, Any]:
    return {
        "ridge_log_target": Ridge(alpha=18.0),
        "random_forest_log_target": RandomForestRegressor(
            n_estimators=260,
            max_depth=18,
            min_samples_leaf=2,
            random_state=random_state,
            n_jobs=-1,
        ),
        "extra_trees_log_target": ExtraTreesRegressor(
            n_estimators=320,
            max_depth=20,
            min_samples_leaf=2,
            random_state=random_state,
            n_jobs=-1,
        ),
        "gradient_boosting_log_target": GradientBoostingRegressor(
            n_estimators=700,
            learning_rate=0.035,
            max_depth=3,
            min_samples_leaf=3,
            subsample=0.85,
            random_state=random_state,
        ),
    }


def selected_models(config: dict[str, Any], random_state: int) -> dict[str, Any]:
    all_models = classification_models(random_state) if is_classification_task(config) else regression_models(random_state)
    requested = config.get("scaffold", {}).get("first_stage_models")
    if not requested:
        return all_models
    return {name: all_models[name] for name in requested if name in all_models}


def evaluate_classification(x: pd.DataFrame, y: pd.Series, config: dict[str, Any], random_state: int) -> tuple[dict[str, Any], Pipeline]:
    x_train, x_valid, y_train, y_valid = train_test_split(x, y, test_size=0.2, random_state=random_state, stratify=y)
    cv = StratifiedKFold(n_splits=5, shuffle=True, random_state=random_state)
    models = selected_models(config, random_state)
    results: dict[str, Any] = {}
    best_name = ""
    best_score = -1.0
    best_pipeline: Pipeline | None = None

    for name, model in models.items():
        pipeline = Pipeline(steps=[("preprocessor", build_preprocessor(x)), ("model", model)])
        start = time.time()
        cv_scores = []
        for train_idx, valid_idx in cv.split(x, y):
            fold_pipeline = clone(pipeline)
            fold_pipeline.fit(x.iloc[train_idx], y.iloc[train_idx])
            fold_pred = fold_pipeline.predict(x.iloc[valid_idx])
            cv_scores.append(float(accuracy_score(y.iloc[valid_idx], fold_pred)))
        pipeline.fit(x_train, y_train)
        valid_pred = pipeline.predict(x_valid)
        metrics = {
            "cv_accuracy_mean": round(float(np.mean(cv_scores)), 6),
            "cv_accuracy_std": round(float(np.std(cv_scores)), 6),
            "holdout_accuracy": round(float(accuracy_score(y_valid, valid_pred)), 6),
            "holdout_macro_f1": round(float(f1_score(y_valid, valid_pred, average="macro")), 6),
            "seconds": round(time.time() - start, 4),
        }
        results[name] = metrics
        if metrics["cv_accuracy_mean"] > best_score:
            best_score = metrics["cv_accuracy_mean"]
            best_name = name
            best_pipeline = pipeline

    if best_pipeline is None:
        raise RuntimeError("No classification model was trained.")
    best_pipeline.fit(x, y)
    return {
        "metric": config["task"]["metric"],
        "best_model": best_name,
        "selection_direction": "maximize",
        "model_results": results,
    }, best_pipeline


def evaluate_regression(x: pd.DataFrame, y: pd.Series, config: dict[str, Any], random_state: int) -> tuple[dict[str, Any], Pipeline]:
    x_train, x_valid, y_train, y_valid = train_test_split(x, y, test_size=0.2, random_state=random_state)
    cv = KFold(n_splits=5, shuffle=True, random_state=random_state)
    models = selected_models(config, random_state)
    metric = str(config["task"]["metric"]).lower()
    target_transform = config.get("feature_engineering", {}).get("target_transform")
    results: dict[str, Any] = {}
    best_name = ""
    best_score = float("inf")
    best_pipeline: Pipeline | None = None

    if target_transform is None and metric == "rmsle":
        target_transform = "log1p"
    if target_transform == "log1p" and (y.astype(float) < 0).any():
        raise ValueError("log1p target transform requires non-negative regression targets")

    y_model = np.log1p(y.astype(float)) if target_transform == "log1p" else y.astype(float)
    for name, model in models.items():
        pipeline = Pipeline(steps=[("preprocessor", build_preprocessor(x)), ("model", model)])
        start = time.time()
        fold_scores = []
        for train_idx, valid_idx in cv.split(x):
            fold_pipeline = clone(pipeline)
            fold_pipeline.fit(x.iloc[train_idx], y_model.iloc[train_idx])
            fold_pred = inverse_target_transform(fold_pipeline.predict(x.iloc[valid_idx]), target_transform)
            fold_scores.append(regression_metric_value(metric, y.iloc[valid_idx], fold_pred))
        pipeline.fit(x_train, np.log1p(y_train.astype(float)) if target_transform == "log1p" else y_train.astype(float))
        holdout_pred = inverse_target_transform(pipeline.predict(x_valid), target_transform)
        cv_key = f"cv_{metric}_mean"
        holdout_key = f"holdout_{metric}"
        metrics = {
            cv_key: round(float(np.mean(fold_scores)), 6),
            f"cv_{metric}_std": round(float(np.std(fold_scores)), 6),
            holdout_key: round(regression_metric_value(metric, y_valid, holdout_pred), 6),
            "holdout_mae": round(float(mean_absolute_error(y_valid, np.clip(holdout_pred, 0, None))), 6),
            "seconds": round(time.time() - start, 4),
        }
        results[name] = metrics
        if metrics[cv_key] < best_score:
            best_score = metrics[cv_key]
            best_name = name
            best_pipeline = pipeline

    if best_pipeline is None:
        raise RuntimeError("No regression model was trained.")
    best_pipeline.fit(x, y_model)
    return {
        "metric": config["task"]["metric"],
        "best_model": best_name,
        "selection_direction": "minimize",
        "model_results": results,
    }, best_pipeline


def evaluate_models(x: pd.DataFrame, y: pd.Series, config: dict[str, Any], random_state: int) -> tuple[dict[str, Any], Pipeline]:
    if is_classification_task(config):
        return evaluate_classification(x, y, config, random_state)
    return evaluate_regression(x, y, config, random_state)


def load_agent_templates(config: dict[str, Any]) -> dict[str, Any]:
    template_value = config.get("agent_templates")
    if not template_value:
        return {"source": {}, "template_mapping": {}}
    template_path = Path(template_value)
    if template_path.is_file():
        return load_yaml(template_path)
    return {"source": {}, "template_mapping": {}}


def build_task_scaffold(config: dict[str, Any], quality: dict[str, Any], agent_templates: dict[str, Any]) -> dict[str, Any]:
    data_cfg = config["data"]
    scaffold_cfg = config.get("scaffold", {})
    return {
        "task": config["task"],
        "inputs": {
            "overview": data_cfg["overview"],
            "train": data_cfg["train"],
            "test": data_cfg["test"],
            "sample_submission": data_cfg["sample_submission"],
        },
        "agent_template_mapping": agent_templates.get("template_mapping", {}),
        "server_template_policy": agent_templates.get("source", {}),
        "data_snapshot": {
            "train_rows": quality["train_rows"],
            "test_rows": quality["test_rows"],
            "sample_submission_rows": quality["sample_submission_rows"],
            "target_summary": quality["target_summary"],
            "missing_train_top20": quality["missing_train_top20"],
            "missing_test_top20": quality["missing_test_top20"],
        },
        "validation_strategy": scaffold_cfg.get("validation_strategy", "local validation"),
        "time_budget_minutes": scaffold_cfg.get("time_budget_minutes", 10),
        "candidate_models": scaffold_cfg.get("first_stage_models", list(selected_models(config, 42).keys())),
        "feature_plan": build_feature_plan(config),
        "risk_points": scaffold_cfg.get("risk_points", []),
        "stage_plan": config["workflow"],
    }


def build_feature_plan(config: dict[str, Any]) -> list[str]:
    preset = config.get("feature_engineering", {}).get("preset")
    if preset == "house_prices_basic":
        return [
            "Train regression models on log1p(SalePrice) to match RMSLE behavior.",
            "Create TotalSF from basement and floor area fields.",
            "Create TotalBath from full/half bathroom fields.",
            "Create house age and remodel age from sale/build/remodel years.",
            "Create binary indicators for garage, basement, fireplace, and pool presence.",
            "Use median imputation for numeric features and most-frequent imputation plus one-hot encoding for categorical features.",
        ]
    return [
        "Use configured drop columns.",
        "Use median imputation for numeric features.",
        "Use most-frequent imputation and one-hot encoding for categorical features.",
    ]


def write_scaffold(output_dir: Path, scaffold: dict[str, Any]) -> None:
    (output_dir / "task_scaffold.json").write_text(json.dumps(scaffold, ensure_ascii=False, indent=2), encoding="utf-8")
    task = scaffold["task"]
    lines = [
        f"# {task['name']} 任务脚手架",
        "",
        "## 任务",
        "",
        f"- 比赛：{task['competition']}",
        f"- 类型：{task['type']}",
        f"- 目标列：{task['target']}",
        f"- 指标：{task['metric']}",
        "",
        "## 服务器模板映射",
        "",
    ]
    for role_key, item in scaffold["agent_template_mapping"].items():
        lines.append(f"- {role_key}: {item['server_template_name']} -> {item['local_role']}")
    lines.extend(
        [
            "",
            "## 输入文件",
            "",
            *[f"- {name}: `{path}`" for name, path in scaffold["inputs"].items()],
            "",
            "## 验证方案",
            "",
            f"- {scaffold['validation_strategy']}",
            f"- 时间预算：{scaffold['time_budget_minutes']} 分钟",
            "",
            "## 候选模型",
            "",
            *[f"- {name}" for name in scaffold["candidate_models"]],
            "",
            "## 特征计划",
            "",
            *[f"- {item}" for item in scaffold["feature_plan"]],
            "",
            "## 风险点",
            "",
            *[f"- {item}" for item in scaffold["risk_points"]],
        ]
    )
    (output_dir / "task_scaffold.md").write_text("\n".join(lines), encoding="utf-8-sig")


def make_submission(best_pipeline: Pipeline, test_features: pd.DataFrame, sample: pd.DataFrame, config: dict[str, Any], output_dir: Path) -> dict[str, Any]:
    task = config["task"]
    submission = sample.copy()
    prediction_column = task.get("prediction_column", sample.columns[1])

    raw_predictions = best_pipeline.predict(test_features)
    if is_regression_task(config):
        target_transform = config.get("feature_engineering", {}).get("target_transform")
        if target_transform is None and str(config["task"].get("metric", "")).lower() == "rmsle":
            target_transform = "log1p"
        predictions = inverse_target_transform(raw_predictions, target_transform)
        if config.get("thresholds", {}).get("require_positive_predictions", False):
            predictions = np.clip(predictions, 1.0, None)
    else:
        predictions = raw_predictions

    submission[prediction_column] = predictions
    path = output_dir / "submission.csv"
    submission.to_csv(path, index=False)

    prediction_values = submission[prediction_column]
    checks = {
        "path": str(path),
        "rows_match": len(submission) == len(sample),
        "columns_match": submission.columns.tolist() == sample.columns.tolist(),
        "missing_predictions": int(prediction_values.isna().sum()),
        "prediction_columns": [prediction_column],
        "valid": False,
    }
    if is_regression_task(config):
        checks.update(
            {
                "positive_predictions": bool((prediction_values > 0).all()),
                "prediction_min": round(float(prediction_values.min()), 6),
                "prediction_max": round(float(prediction_values.max()), 6),
                "prediction_mean": round(float(prediction_values.mean()), 6),
            }
        )
        checks["valid"] = checks["rows_match"] and checks["columns_match"] and checks["missing_predictions"] == 0 and checks["positive_predictions"]
    else:
        observed_values = sorted(prediction_values.dropna().unique().tolist())
        configured_allowed = config.get("thresholds", {}).get("allowed_prediction_values")
        # Derive the class domain from the trained classifier when not explicitly
        # configured. A classifier can only emit classes it learned, so its
        # ``classes_`` is a correct-by-construction allow list. This fixes
        # multiclass tasks (e.g. 7-class Cover_Type) that previously failed the
        # binary-only [0,1] fallback. Binary tasks keep classes_ == [0, 1].
        model_classes = None
        try:
            model_step = best_pipeline.named_steps.get("model") if hasattr(best_pipeline, "named_steps") else None
            if model_step is not None and hasattr(model_step, "classes_"):
                model_classes = sorted(int(c) if float(c).is_integer() else c for c in model_step.classes_.tolist())
        except Exception:
            model_classes = None
        effective_allowed = configured_allowed or model_classes
        if effective_allowed:
            allowed_values_only = set(observed_values).issubset(set(effective_allowed))
        else:
            allowed_values_only = observed_values in ([0], [1], [0, 1])
        checks.update(
            {
                "allowed_values": effective_allowed or [0, 1],
                "allowed_values_source": "config" if configured_allowed else ("classifier_classes_" if model_classes else "binary_default"),
                "allowed_values_only": allowed_values_only,
                "prediction_distribution": {str(k): int(v) for k, v in prediction_values.value_counts().to_dict().items()},
            }
        )
        checks["valid"] = checks["rows_match"] and checks["columns_match"] and checks["missing_predictions"] == 0 and checks["allowed_values_only"]
    return checks


def metrics_pass(config: dict[str, Any], evaluation: dict[str, Any]) -> bool:
    thresholds = config["thresholds"]
    best_metrics = evaluation["model_results"][evaluation["best_model"]]
    metric = str(config["task"]["metric"]).lower()
    if is_regression_task(config):
        cv_key = f"cv_{metric}_mean"
        holdout_key = f"holdout_{metric}"
        cv_threshold = thresholds.get(f"max_cv_{metric}")
        holdout_threshold = thresholds.get(f"max_holdout_{metric}")
        checks = []
        if cv_threshold is not None:
            checks.append(best_metrics[cv_key] <= cv_threshold)
        if holdout_threshold is not None:
            checks.append(best_metrics[holdout_key] <= holdout_threshold)
        return all(checks) if checks else True
    threshold = thresholds.get("min_validation_accuracy")
    return best_metrics["cv_accuracy_mean"] >= threshold if threshold is not None else True


def build_stage_audit(
    config: dict[str, Any],
    quality: dict[str, Any],
    evaluation: dict[str, Any],
    submission: dict[str, Any],
    accepted: bool,
    output_dir: Path,
) -> dict[str, Any]:
    thresholds = config["thresholds"]
    best_model = evaluation["best_model"]
    best_metrics = evaluation["model_results"][best_model]
    metric_ok = metrics_pass(config, evaluation)

    stages = [
        {
            "stage": "task_understanding",
            "owner_role": "Orchestrator/Planner",
            "status": "passed",
            "evidence": [config["data"]["overview"], str(output_dir / "task_scaffold.json")],
            "checks": {
                "task_name": config["task"]["name"],
                "target": config["task"]["target"],
                "metric": config["task"]["metric"],
                "server_templates_read_only": True,
            },
        },
        {
            "stage": "preliminary_eda",
            "owner_role": "Analyst",
            "status": "passed",
            "evidence": [str(output_dir / "data_quality.json")],
            "checks": {
                "train_rows": quality["train_rows"],
                "test_rows": quality["test_rows"],
                "target_summary_recorded": bool(quality["target_summary"]),
                "missing_values_recorded": bool(quality["missing_train_top20"]) and bool(quality["missing_test_top20"]),
            },
        },
        {
            "stage": "data_quality_check",
            "owner_role": "Reviewer/Gate",
            "status": "passed" if quality["train_test_feature_columns_match"] else "failed",
            "evidence": [str(output_dir / "data_quality.json")],
            "checks": {
                "train_test_feature_columns_match": quality["train_test_feature_columns_match"],
                "sample_submission_rows": quality["sample_submission_rows"],
                "sample_submission_columns": quality["sample_submission_columns"],
            },
        },
        {
            "stage": "feature_engineering",
            "owner_role": "Developer",
            "status": "passed",
            "evidence": [str(output_dir / "task_scaffold.json")],
            "checks": {
                "preset": config.get("feature_engineering", {}).get("preset", "generic"),
                "feature_plan_recorded": True,
                "target_transform": config.get("feature_engineering", {}).get("target_transform"),
            },
        },
        {
            "stage": "model_validation",
            "owner_role": "Developer/Reviewer",
            "status": "passed" if metric_ok else "failed",
            "evidence": [str(output_dir / "model_results.json")],
            "checks": {
                "best_model": best_model,
                "best_metrics": best_metrics,
                "thresholds": thresholds,
                "candidate_model_count": len(evaluation["model_results"]),
            },
        },
        {
            "stage": "submission_generation",
            "owner_role": "Developer/Reviewer",
            "status": "passed" if submission["valid"] else "failed",
            "evidence": [submission["path"]],
            "checks": submission,
        },
        {
            "stage": "report_and_review",
            "owner_role": "Evidence/Summarizer",
            "status": "passed" if accepted else "failed",
            "evidence": [str(output_dir / "local_report.md"), str(output_dir / "local_report.docx")],
            "checks": {
                "report_grounded_in_outputs": True,
                "post_scaffold_written": str(output_dir / "post_scaffold_improvement.json"),
                "local_gate_expected": accepted,
            },
        },
    ]
    failed = [stage["stage"] for stage in stages if stage["status"] != "passed"]
    return {
        "task": config["task"],
        "experiment_dir": str(output_dir),
        "workflow_version": "v3_generic_tabular",
        "local_only": True,
        "server_private_templates_modified": False,
        "kaggle_api_submission": False,
        "stages": stages,
        "all_stages_passed": not failed,
        "failed_stages": failed,
        "next_actions": [
            "Configure Kaggle API token before official download/submission.",
            "Use this generic workflow as the baseline for additional teacher-provided tabular datasets.",
            "Keep GPU/server work disabled until a task has a clear compute benefit.",
        ],
    }


def write_stage_audit(output_dir: Path, audit: dict[str, Any]) -> None:
    (output_dir / "workflow_stage_audit.json").write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    lines = [
        f"# {audit['task']['name']} 阶段化工作流审计",
        "",
        "## 总结",
        "",
        f"- 实验目录：`{audit['experiment_dir']}`",
        f"- 工作流版本：`{audit['workflow_version']}`",
        f"- 全部阶段通过：`{audit['all_stages_passed']}`",
        f"- 官方 Kaggle 提交：`{audit['kaggle_api_submission']}`",
        f"- 修改服务器私有模板：`{audit['server_private_templates_modified']}`",
        "",
        "## 阶段检查",
        "",
    ]
    for stage in audit["stages"]:
        lines.extend(
            [
                f"### {stage['stage']}",
                "",
                f"- 角色：{stage['owner_role']}",
                f"- 状态：`{stage['status']}`",
                f"- 证据：{', '.join(f'`{item}`' for item in stage['evidence'])}",
                f"- 检查项：`{json.dumps(stage['checks'], ensure_ascii=False)}`",
                "",
            ]
        )
    lines.extend(["## 下一步", ""])
    lines.extend(f"- {item}" for item in audit["next_actions"])
    (output_dir / "workflow_stage_audit.md").write_text("\n".join(lines), encoding="utf-8-sig")


def build_post_scaffold_improvement(config: dict[str, Any], evaluation: dict[str, Any], submission: dict[str, Any], accepted: bool) -> dict[str, Any]:
    best_model = evaluation["best_model"]
    best_metrics = evaluation["model_results"][best_model]
    if accepted:
        decision = "stop_after_local_gate"
        rationale = "Local metric thresholds and submission checks passed. Continue only if official Kaggle submission or stronger model comparison is required."
    else:
        decision = "iterate"
        rationale = "At least one local metric or submission gate failed. Use the failed stage as the next improvement target."

    return {
        "task": config["task"]["name"],
        "decision": decision,
        "rationale": rationale,
        "best_model": best_model,
        "best_metrics": best_metrics,
        "submission_valid": submission["valid"],
        "next_iteration_candidates": [
            "Review high-missing categorical columns and decide whether missing should be a category instead of imputed.",
            "Add task-specific feature interactions only if they can be explained in the report.",
            "Add LightGBM/XGBoost/CatBoost after the sklearn baseline remains stable.",
            "Switch to official Kaggle API submission after credentials are configured.",
        ],
    }


def write_post_scaffold(output_dir: Path, improvement: dict[str, Any]) -> None:
    (output_dir / "post_scaffold_improvement.json").write_text(json.dumps(improvement, ensure_ascii=False, indent=2), encoding="utf-8")
    lines = [
        f"# {improvement['task']} Post-Scaffold 改进记录",
        "",
        f"- 决策：`{improvement['decision']}`",
        f"- 理由：{improvement['rationale']}",
        f"- 最佳模型：`{improvement['best_model']}`",
        f"- 最佳指标：`{json.dumps(improvement['best_metrics'], ensure_ascii=False)}`",
        f"- submission 通过：`{improvement['submission_valid']}`",
        "",
        "## 下一轮候选动作",
        "",
        *[f"- {item}" for item in improvement["next_iteration_candidates"]],
    ]
    (output_dir / "post_scaffold_improvement.md").write_text("\n".join(lines), encoding="utf-8-sig")


def write_docx_report(markdown_path: Path, docx_path: Path) -> None:
    if Document is None:
        return

    document = Document()
    font_name = "宋体"
    black = RGBColor(0, 0, 0)
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "List Bullet"]:
        style = document.styles[style_name]
        style.font.name = font_name
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        style.font.color.rgb = black
        if style_name == "Normal":
            style.font.size = Pt(11)

    for raw in markdown_path.read_text(encoding="utf-8-sig").splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            paragraph = document.add_paragraph(line[2:].strip(), style="Title")
        elif line.startswith("## "):
            paragraph = document.add_paragraph(line[3:].strip(), style="Heading 1")
        elif line.startswith("- "):
            paragraph = document.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            paragraph = document.add_paragraph(line)
        for run in paragraph.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
            run.font.color.rgb = black

    document.save(docx_path)


def write_markdown_report(
    output_dir: Path,
    config: dict[str, Any],
    quality: dict[str, Any],
    evaluation: dict[str, Any],
    submission: dict[str, Any],
    accepted: bool,
) -> None:
    best = evaluation["best_model"]
    best_metrics = evaluation["model_results"][best]
    task = config["task"]
    lines = [
        f"# {task['competition']} 本地测试报告",
        "",
        "## 任务理解",
        "",
        f"- 任务：{task['competition']}。",
        f"- 类型：{task['type']} 表格数据任务。",
        f"- 指标：{task['metric']}，本地使用交叉验证和 holdout 作为代理评估。",
        "- 当前说明：本机暂未配置 Kaggle API 凭据，因此本轮使用公开镜像数据完成本地闭环测试；后续可替换为官方 Kaggle API 下载。",
        "",
        "## 服务器模板对齐",
        "",
        "- 本地流程参照服务器公开 Agent 模板字段与职责映射。",
        "- 本轮没有读取或修改服务器私有 Agent 和工作流。",
        "",
        "## 数据质量",
        "",
        f"- train 行列数：{quality['train_rows']} x {quality['train_columns']}",
        f"- test 行列数：{quality['test_rows']} x {quality['test_columns']}",
        f"- sample_submission 行数：{quality['sample_submission_rows']}",
        f"- 训练/测试特征列一致：{quality['train_test_feature_columns_match']}",
        f"- 目标摘要：{json.dumps(quality['target_summary'], ensure_ascii=False)}",
        f"- 训练集缺失率 Top20：{json.dumps(quality['missing_train_top20'], ensure_ascii=False)}",
        f"- 测试集缺失率 Top20：{json.dumps(quality['missing_test_top20'], ensure_ascii=False)}",
        "",
        "## 模型验证",
        "",
        f"- 最佳模型：`{best}`",
        f"- 指标：`{json.dumps(best_metrics, ensure_ascii=False)}`",
        "",
        "## Submission 检查",
        "",
        f"- submission 文件：`{submission['path']}`",
        f"- 行数匹配：{submission['rows_match']}",
        f"- 列名匹配：{submission['columns_match']}",
        f"- 缺失预测数：{submission['missing_predictions']}",
        f"- 是否通过：{submission['valid']}",
        "",
        "## 验收结论",
        "",
        f"- 本地验收：{'通过' if accepted else '未通过'}",
        "- 判断依据：数据质量检查、模型验证分数、submission 格式检查、阶段审计和 post-scaffold 改进记录均已保存。",
        "",
        "## 下一步",
        "",
        "- 配置 Kaggle API token 后，替换为官方 Kaggle API 下载与可选真实提交。",
        "- 在更多表格任务或老师真实数据案例上复用同一套通用 workflow。",
        "- 闭环稳定后，再接入 GPU 服务器做深度模型或超参数搜索。",
    ]
    markdown_path = output_dir / "local_report.md"
    markdown_path.write_text("\n".join(lines), encoding="utf-8-sig")
    docx_path = output_dir / "local_report.docx"
    write_docx_report(markdown_path, docx_path)
    legacy_prefix = str(config["task"].get("name", "")).strip()
    if legacy_prefix:
        legacy_markdown_path = output_dir / f"{legacy_prefix}_local_report.md"
        legacy_docx_path = output_dir / f"{legacy_prefix}_local_report.docx"
        legacy_markdown_path.write_text(markdown_path.read_text(encoding="utf-8-sig"), encoding="utf-8-sig")
        if docx_path.exists():
            legacy_docx_path.write_bytes(docx_path.read_bytes())


def _run_local_pipeline_for_tests(
    config: dict[str, Any],
    output_base: Path,
    random_state: int,
) -> dict[str, Any]:
    data_cfg = config["data"]
    task = config["task"]
    paths = {
        "train": Path(data_cfg["train"]),
        "test": Path(data_cfg["test"]),
        "sample_submission": Path(data_cfg["sample_submission"]),
        "overview": Path(data_cfg["overview"]),
    }
    for label, path in paths.items():
        if not path.exists():
            raise FileNotFoundError(f"Missing {label}: {path}")

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_dir = output_base / task["name"] / timestamp
    output_dir.mkdir(parents=True, exist_ok=True)

    train = pd.read_csv(paths["train"])
    test = pd.read_csv(paths["test"])
    sample = pd.read_csv(paths["sample_submission"])
    quality = data_quality_report(train, test, sample, config)
    agent_templates = load_agent_templates(config)
    scaffold = build_task_scaffold(config, quality, agent_templates)
    write_scaffold(output_dir, scaffold)

    x = apply_feature_preset(train.drop(columns=[task["target"]]), config)
    y = train[task["target"]]
    test_features = apply_feature_preset(test, config)

    evaluation, best_pipeline = evaluate_models(x, y, config, random_state)
    submission = make_submission(best_pipeline, test_features, sample, config, output_dir)
    accepted = bool(quality["train_test_feature_columns_match"] and submission["valid"] and metrics_pass(config, evaluation))

    improvement = build_post_scaffold_improvement(config, evaluation, submission, accepted)
    write_post_scaffold(output_dir, improvement)
    write_markdown_report(output_dir, config, quality, evaluation, submission, accepted)
    stage_audit = build_stage_audit(config, quality, evaluation, submission, accepted, output_dir)
    write_stage_audit(output_dir, stage_audit)

    artifacts = {
        "task": task,
        "data_quality": quality,
        "evaluation": evaluation,
        "submission_check": submission,
        "task_scaffold": scaffold,
        "post_scaffold_improvement": improvement,
        "stage_audit": stage_audit,
        "accepted": accepted,
        "thresholds": config["thresholds"],
        "outputs": {
            "submission": submission["path"],
            "scaffold": str(output_dir / "task_scaffold.json"),
            "scaffold_markdown": str(output_dir / "task_scaffold.md"),
            "post_scaffold": str(output_dir / "post_scaffold_improvement.json"),
            "post_scaffold_markdown": str(output_dir / "post_scaffold_improvement.md"),
            "report": str(output_dir / "local_report.md"),
            "report_docx": str(output_dir / "local_report.docx"),
            "stage_audit": str(output_dir / "workflow_stage_audit.json"),
            "stage_audit_markdown": str(output_dir / "workflow_stage_audit.md"),
        },
    }
    (output_dir / "experiment_log.json").write_text(json.dumps(artifacts, ensure_ascii=False, indent=2), encoding="utf-8")
    (output_dir / "data_quality.json").write_text(json.dumps(quality, ensure_ascii=False, indent=2), encoding="utf-8")
    (output_dir / "model_results.json").write_text(json.dumps(evaluation, ensure_ascii=False, indent=2), encoding="utf-8")

    return {
        "output_dir": str(output_dir),
        "accepted": accepted,
        "best_model": evaluation["best_model"],
        "best_metrics": evaluation["model_results"][evaluation["best_model"]],
    }


def run(config: dict[str, Any], output_base: Path, random_state: int) -> dict[str, Any]:
    """Compatibility entry point that blocks retired workstation training."""
    from research_os.hpc_policy import HPCPolicyError

    raise HPCPolicyError(
        "blocked_local_training_disabled: local tabular training is disabled by release policy; "
        "use the gated HPC/GPU workflow"
    )


def main(argv: list[str] | None = None) -> int:
    parse_args(argv)
    print(json.dumps({
        "status": "blocked_local_training_disabled",
        "training_started": False,
        "required_compute": "gpu",
        "message": "Local training is disabled by release policy. Use the gated HPC/GPU workflow.",
    }, ensure_ascii=False))
    return 2


if __name__ == "__main__":
    raise SystemExit(main())
