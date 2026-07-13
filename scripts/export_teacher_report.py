from __future__ import annotations

import json
import math
import re
import textwrap
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
SOURCE_MD = DOCS_DIR / "AI-Research-Workstation-Multi-Agent-Framework-Report-20260614.md"
ASSET_DIR = DOCS_DIR / "report_assets_20260615"

DOCX_PATH = DOCS_DIR / "AI科研工作站多Agent框架汇报-20260615.docx"
TEX_PATH = DOCS_DIR / "AI-Research-Workstation-Teacher-Report-20260615.tex"
SUMMARY_PATH = DOCS_DIR / "AI-Research-Workstation-Teacher-Report-20260615.export.json"

EAST_ASIA_FONT = "Noto Sans SC"
LATIN_FONT = "Arial"
MONO_FONT = "Consolas"


def font_path(*names: str) -> str:
    font_dir = Path("C:/Windows/Fonts")
    for name in names:
        candidate = font_dir / name
        if candidate.exists():
            return str(candidate)
    return str(font_dir / "simhei.ttf")


FONT_REGULAR_PATH = font_path("NotoSansSC-VF.ttf", "msyh.ttc", "simhei.ttf", "Deng.ttf")
FONT_BOLD_PATH = font_path("simhei.ttf", "Dengb.ttf", "NotoSansSC-VF.ttf")


def pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(FONT_BOLD_PATH if bold else FONT_REGULAR_PATH, size)


def wrap_for_box(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    lines: list[str] = []
    for raw in text.split("\n"):
        current = ""
        for ch in raw:
            trial = current + ch
            if draw.textbbox((0, 0), trial, font=font)[2] <= max_width or not current:
                current = trial
            else:
                lines.append(current)
                current = ch
        if current:
            lines.append(current)
    return lines


def draw_text_box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    text: str,
    *,
    fill: str,
    outline: str = "#2F3A4A",
    text_fill: str = "#102033",
    radius: int = 18,
    title: bool = False,
    font_size: int | None = None,
) -> None:
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=3)
    font = pil_font(font_size if font_size is not None else (40 if title else 34), bold=title)
    lines = wrap_for_box(draw, text, font, x2 - x1 - 42)
    line_heights = [draw.textbbox((0, 0), line, font=font)[3] for line in lines]
    total_h = sum(line_heights) + max(0, len(lines) - 1) * 10
    y = y1 + ((y2 - y1 - total_h) // 2)
    for idx, line in enumerate(lines):
        bbox = draw.textbbox((0, 0), line, font=font)
        x = x1 + ((x2 - x1 - (bbox[2] - bbox[0])) // 2)
        draw.text((x, y), line, font=font, fill=text_fill)
        y += line_heights[idx] + 10


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#506176") -> None:
    draw.line([start, end], fill=color, width=4)
    x1, y1 = start
    x2, y2 = end
    angle = math.atan2(y2 - y1, x2 - x1)
    arrow_len = 18
    arrow_angle = math.radians(28)
    points = [
        (x2, y2),
        (
            int(x2 - arrow_len * math.cos(angle - arrow_angle)),
            int(y2 - arrow_len * math.sin(angle - arrow_angle)),
        ),
        (
            int(x2 - arrow_len * math.cos(angle + arrow_angle)),
            int(y2 - arrow_len * math.sin(angle + arrow_angle)),
        ),
    ]
    draw.polygon(points, fill=color)


def draw_polyline_arrow(draw: ImageDraw.ImageDraw, points: list[tuple[int, int]], color: str = "#506176") -> None:
    if len(points) < 2:
        return
    draw.line(points, fill=color, width=4, joint="curve")
    draw_arrow(draw, points[-2], points[-1], color)


def save_diagram(path: Path, title: str, subtitle: str, draw_body) -> None:
    image = Image.new("RGB", (1800, 1050), "#F6F8FB")
    draw = ImageDraw.Draw(image)
    draw.text((70, 45), title, font=pil_font(48, bold=True), fill="#122033")
    draw.text((72, 112), subtitle, font=pil_font(25), fill="#526174")
    draw.line((70, 160, 1730, 160), fill="#D8DEE8", width=3)
    draw_body(draw)
    image.save(path, quality=95)


def create_diagrams() -> list[dict[str, str]]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    diagrams: list[dict[str, str]] = []

    roles_path = ASSET_DIR / "01_multi_agent_roles.png"

    def roles(draw: ImageDraw.ImageDraw) -> None:
        draw_text_box(
            draw,
            (80, 395, 410, 585),
            "总控 Agent\n任务分解 / 调度\n证据汇总 / 回滚",
            fill="#DDEBFF",
            outline="#3971B8",
            title=True,
            font_size=30,
        )
        role_labels = [
            "研究背景 Agent\n文献 / 公开资料",
            "数据审计 Agent\n版本 / 缺失 / 泄露",
            "特征工程 Agent\n领域特征 / 消融",
            "模型选择 Agent\n路线 / 参数空间",
            "代码实现 Agent\n训练脚本 / Runner",
            "HPC/GPU 执行 Agent\n远程训练 / 日志回传",
            "验证分析 Agent\nCV / OOF / 误差分析",
            "提交门禁 Agent\nSchema / 风险 / 授权",
            "报告总结 Agent\n台账 / 阶段汇报",
            "反思审查 Agent\n逻辑 / 复现 / 过拟合",
        ]
        colors = ["#E8F4EC", "#FFF3D8", "#F2EAFB", "#EAF6F9", "#FDE9E7"]
        x0, y0, w, h = 500, 210, 560, 110
        row_gap = 132
        positions: list[tuple[int, int, int, int]] = []
        for idx, label in enumerate(role_labels):
            col = idx % 2
            row = idx // 2
            x = x0 + col * 615
            y = y0 + row * row_gap
            box = (x, y, x + w, y + h)
            positions.append(box)
            draw_text_box(draw, box, label, fill=colors[idx % len(colors)], radius=14)
        bus_x = 455
        draw.line((bus_x, y0 + h // 2, bus_x, y0 + 4 * row_gap + h // 2), fill="#506176", width=4)
        draw_arrow(draw, (410, 490), (bus_x, 490))
        for row in range(5):
            left = positions[row * 2]
            right = positions[row * 2 + 1]
            y_mid = left[1] + h // 2
            draw_arrow(draw, (bus_x, y_mid), (left[0], y_mid))
            draw_arrow(draw, (left[2], y_mid), (right[0], y_mid))
        draw_text_box(
            draw,
            (80, 900, 1720, 985),
            "关键思想：每个 Agent 只接收必要上下文，并交付可检查 artifact；失败时只回滚对应模块。",
            fill="#FFFFFF",
            outline="#D8DEE8",
            text_fill="#233143",
            radius=12,
        )

    save_diagram(
        roles_path,
        "多 Agent 角色分工框架",
        "用职责边界和上下文隔离，缓解单模型长流程科研任务中的上下文超载",
        roles,
    )
    diagrams.append({"path": str(roles_path), "caption": "图 1  多 Agent 角色分工框架"})

    architecture_path = ASSET_DIR / "02_system_architecture.png"

    def architecture(draw: ImageDraw.ImageDraw) -> None:
        layers = [
            ("输入层", ["Kaggle 官方数据", "论文 / 文档 / Discussion", "历史实验台账", "老师反馈 / 研究目标"], "#E8F0FE"),
            ("多 Agent 编排层", ["任务拆解", "上下文切片", "模型-角色匹配", "失败回滚"], "#E9F8F0"),
            ("执行层", ["本地代码生成", "HPC/GPU 训练", "Artifact 回传", "安全密钥管理"], "#FFF4DC"),
            ("验证层", ["Stratified KFold", "OOF 预测", "误差分析", "多模型融合"], "#F1EAFE"),
            ("治理层", ["实验日志", "Submission 门禁", "复现报告", "人工确认"], "#FDECEA"),
        ]
        col_w, gap, top, bottom = 300, 35, 230, 850
        start_x = 70
        for idx, (name, items, color) in enumerate(layers):
            x = start_x + idx * (col_w + gap)
            draw_text_box(draw, (x, top, x + col_w, top + 90), name, fill=color, title=True, radius=16)
            for item_idx, item in enumerate(items):
                y = top + 125 + item_idx * 96
                draw_text_box(draw, (x + 22, y, x + col_w - 22, y + 70), item, fill="#FFFFFF", outline="#B8C1CF", radius=12)
            if idx < len(layers) - 1:
                draw_arrow(draw, (x + col_w, (top + bottom) // 2), (x + col_w + gap, (top + bottom) // 2), "#69788A")
        draw.text((75, 925), "架构重点：不把科研过程压进单轮对话，而是沉淀为数据、代码、日志、指标、提交和报告的证据链。", font=pil_font(28, bold=True), fill="#233143")

    save_diagram(
        architecture_path,
        "AI 科研工作站总体架构",
        "从输入、编排、执行、验证到治理，形成可复现的机器学习科研闭环",
        architecture,
    )
    diagrams.append({"path": str(architecture_path), "caption": "图 2  AI 科研工作站总体架构"})

    workflow_path = ASSET_DIR / "03_experiment_workflow.png"

    def workflow(draw: ImageDraw.ImageDraw) -> None:
        steps = [
            ("研究目标\nKaggle 任务", 90, 235),
            ("公开资料\n方法调研", 430, 235),
            ("数据审计\nhash / 泄露", 770, 235),
            ("Baseline\n建立", 1110, 235),
            ("可信 CV 层\nKFold + 多 seed", 1450, 235),
            ("多模型实验\nLGB/XGB/NN", 1450, 505),
            ("OOF 与\n误差分析", 1110, 505),
            ("融合 / 校准\n消融", 770, 505),
            ("提交门禁\n证据是否更优", 430, 505),
            ("人工确认\nKaggle 提交", 90, 505),
            ("榜单反馈\n更新台账", 90, 745),
            ("回滚到对应模块\n继续迭代", 770, 745),
        ]
        box_w, box_h = 250, 100
        boxes = []
        for label, x, y in steps:
            box = (x, y, x + box_w, y + box_h)
            boxes.append(box)
            draw_text_box(draw, box, label, fill="#FFFFFF", outline="#5B7AA5", radius=16)

        def right_mid(idx: int) -> tuple[int, int]:
            x1, y1, x2, y2 = boxes[idx]
            return x2, (y1 + y2) // 2

        def left_mid(idx: int) -> tuple[int, int]:
            x1, y1, x2, y2 = boxes[idx]
            return x1, (y1 + y2) // 2

        def top_mid(idx: int) -> tuple[int, int]:
            x1, y1, x2, _ = boxes[idx]
            return (x1 + x2) // 2, y1

        def bottom_mid(idx: int) -> tuple[int, int]:
            x1, _, x2, y2 = boxes[idx]
            return (x1 + x2) // 2, y2

        for a, b in [(0, 1), (1, 2), (2, 3), (3, 4)]:
            draw_arrow(draw, right_mid(a), left_mid(b), "#526D8D")
        draw_arrow(draw, bottom_mid(4), top_mid(5), "#526D8D")
        for a, b in [(5, 6), (6, 7), (7, 8), (8, 9)]:
            draw_arrow(draw, left_mid(a), right_mid(b), "#526D8D")
        draw_arrow(draw, bottom_mid(9), top_mid(10), "#526D8D")
        draw_polyline_arrow(draw, [left_mid(10), (45, 795), (45, 285), left_mid(0)], "#7D8CA1")

        draw_polyline_arrow(draw, [bottom_mid(8), (555, 690), (895, 690), top_mid(11)], "#C05F4C")
        draw_polyline_arrow(draw, [right_mid(11), (1590, 795), (1590, 640), bottom_mid(5)], "#C05F4C")
        draw.text((560, 645), "门禁不通过：只回滚特征 / 模型 / 参数等相关模块", font=pil_font(26, bold=True), fill="#A04436")
        draw_text_box(
            draw,
            (420, 910, 1380, 990),
            "闭环标准：实验假设 -> 代码实现 -> HPC/GPU 训练 -> CV/OOF -> 决策 -> 台账记录",
            fill="#FFFFFF",
            outline="#D8DEE8",
            text_fill="#233143",
            radius=12,
        )

    save_diagram(
        workflow_path,
        "科研实验闭环流程",
        "以本地 CV、OOF、误差分析和提交门禁决定下一轮实验，而不是只追逐单次榜单反馈",
        workflow,
    )
    diagrams.append({"path": str(workflow_path), "caption": "图 3  科研实验闭环流程"})

    return diagrams


def set_run_font(run, *, size: float | None = None, bold: bool | None = None, color: str | None = None, mono: bool = False) -> None:
    run.font.name = MONO_FONT if mono else LATIN_FONT
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    east_font = MONO_FONT if mono else EAST_ASIA_FONT
    r_fonts.set(qn("w:eastAsia"), east_font)
    r_fonts.set(qn("w:ascii"), MONO_FONT if mono else LATIN_FONT)
    r_fonts.set(qn("w:hAnsi"), MONO_FONT if mono else LATIN_FONT)


def set_style_font(style, *, size: float, bold: bool = False, color: str = "1F2937") -> None:
    style.font.name = LATIN_FONT
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    r_fonts.set(qn("w:ascii"), LATIN_FONT)
    r_fonts.set(qn("w:hAnsi"), LATIN_FONT)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, fill: str | None = None) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade_cell(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    add_inline_runs(paragraph, text, bold=bold, size=9.2)
    if not run.text:
        run._element.getparent().remove(run._element)


def add_inline_runs(paragraph, text: str, *, bold: bool = False, size: float = 10.5, color: str = "1F2937") -> None:
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    for part in pattern.split(text):
        if not part:
            continue
        is_code = part.startswith("`") and part.endswith("`")
        is_bold = part.startswith("**") and part.endswith("**")
        clean = part[1:-1] if is_code else part[2:-2] if is_bold else part
        run = paragraph.add_run(clean)
        set_run_font(run, size=size, bold=bold or is_bold, color=color, mono=is_code)


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.line_spacing = 1.35
    paragraph.paragraph_format.space_after = Pt(5)
    add_inline_runs(paragraph, text.strip())


def add_quote(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, "EAF2FF")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.left_indent = Cm(0.2)
    paragraph.paragraph_format.right_indent = Cm(0.2)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    add_inline_runs(paragraph, text.strip(), bold=True, size=10.5, color="244A7F")


def add_code_block(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, "F3F4F6")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    for idx, line in enumerate(code.splitlines()):
        if idx:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, size=9.2, mono=True, color="374151")


def parse_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        stripped = line.strip()
        if re.fullmatch(r"\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?", stripped):
            continue
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        rows.append(cells)
    return rows


def add_markdown_table(doc: Document, lines: list[str]) -> None:
    rows = parse_table(lines)
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            text = row[c_idx] if c_idx < len(row) else ""
            set_cell_text(table.cell(r_idx, c_idx), text, bold=(r_idx == 0), fill="DDEBFF" if r_idx == 0 else None)
    doc.add_paragraph()


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.9)
    section.left_margin = Cm(2.15)
    section.right_margin = Cm(2.15)
    section.start_type = WD_SECTION_START.NEW_PAGE

    set_style_font(doc.styles["Normal"], size=10.5)
    set_style_font(doc.styles["Heading 1"], size=18, bold=True, color="14345A")
    set_style_font(doc.styles["Heading 2"], size=15, bold=True, color="1D4E89")
    set_style_font(doc.styles["Heading 3"], size=12.5, bold=True, color="2F5F73")
    for style_name in ("List Bullet", "List Number"):
        if style_name in doc.styles:
            set_style_font(doc.styles[style_name], size=10.2)


def add_cover(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("基于多 Agent 协同的 AI 科研工作站框架汇报")
    set_run_font(run, size=24, bold=True, color="14345A")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("面向 Kaggle/HPC 机器学习科研任务的上下文分治与科研闭环系统")
    set_run_font(run, size=13, bold=True, color="536173")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"阶段汇报文件 | PDF + Word | {date.today().isoformat()}")
    set_run_font(run, size=10.5, color="6B7280")

    doc.add_paragraph()
    add_quote(
        doc,
        "核心思想：不是让单个最强模型硬撑完整科研流程，而是把研究、数据、模型、代码、HPC 训练、验证、提交和报告拆成可验证、可回滚、可复现的多 Agent 工作流。",
    )
    doc.add_paragraph()

    overview = doc.add_paragraph()
    overview.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_inline_runs(overview, "汇报重点", bold=True, size=12.5, color="14345A")
    for item in [
        "为什么单模型 1M 上下文仍不足以稳定完成完整 Kaggle 科研流程",
        "多 Agent 上下文分治如何降低遗忘、职责混杂和不可复现风险",
        "系统如何借鉴 MLE-bench，并进一步落到可长期运行的科研工作站",
        "当前 EXP000-EXP005 已形成的数据、训练、提交和报告闭环证据",
    ]:
        paragraph = doc.add_paragraph(style="List Bullet")
        add_inline_runs(paragraph, item, size=10.5)

    doc.add_page_break()


def build_docx(text: str, diagrams: list[dict[str, str]]) -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    diagram_idx = 0
    code_lang: str | None = None
    code_buffer: list[str] = []
    paragraph_buffer: list[str] = []
    table_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            add_body_paragraph(doc, " ".join(part.strip() for part in paragraph_buffer if part.strip()))
            paragraph_buffer = []

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            add_markdown_table(doc, table_buffer)
            table_buffer = []

    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if code_lang is not None:
            if stripped.startswith("```"):
                if code_lang == "mermaid":
                    if diagram_idx < len(diagrams):
                        info = diagrams[diagram_idx]
                        paragraph = doc.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run(info["caption"])
                        set_run_font(run, size=9.5, bold=True, color="374151")
                        doc.add_picture(info["path"], width=Cm(16.4))
                        last = doc.paragraphs[-1]
                        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        diagram_idx += 1
                else:
                    add_code_block(doc, "\n".join(code_buffer))
                code_lang = None
                code_buffer = []
            else:
                code_buffer.append(line)
            i += 1
            continue

        if stripped.startswith("```"):
            flush_paragraph()
            flush_table()
            code_lang = stripped.strip("`").strip().lower() or "text"
            code_buffer = []
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            flush_paragraph()
            table_buffer.append(line)
            i += 1
            continue
        flush_table()

        if not stripped:
            flush_paragraph()
            i += 1
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            level = len(heading.group(1))
            title = heading.group(2).strip()
            if level == 1:
                if title.startswith("基于多 Agent"):
                    i += 1
                    continue
                doc.add_heading(title, level=1)
            elif level == 2:
                doc.add_heading(title, level=1)
            else:
                doc.add_heading(title, level=2)
            i += 1
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            add_quote(doc, stripped.lstrip("> ").strip())
            i += 1
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet or numbered:
            flush_paragraph()
            paragraph = doc.add_paragraph(style="List Bullet" if bullet else "List Number")
            add_inline_runs(paragraph, (bullet or numbered).group(1), size=10.2)
            i += 1
            continue

        paragraph_buffer.append(stripped)
        i += 1

    flush_table()
    flush_paragraph()
    doc.save(DOCX_PATH)


def latex_escape(text: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    return "".join(replacements.get(ch, ch) for ch in text)


def latex_inline(text: str) -> str:
    result: list[str] = []
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    for part in pattern.split(text):
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            result.append(r"\texttt{" + latex_escape(part[1:-1]) + "}")
        elif part.startswith("**") and part.endswith("**"):
            result.append(r"\textbf{" + latex_escape(part[2:-2]) + "}")
        else:
            result.append(latex_escape(part))
    return "".join(result)


def strip_heading_number(title: str) -> str:
    return re.sub(r"^\d+(?:\.\d+)*\s*", "", title).strip()


def latex_caption(caption: str) -> str:
    return re.sub(r"^图\s*\d+\s*", "", caption).strip()


def latex_table(lines: list[str]) -> str:
    rows = parse_table(lines)
    if not rows:
        return ""
    cols = max(len(row) for row in rows)
    width = 0.78 / max(cols, 1)
    spec = "".join([rf">{{\raggedright\arraybackslash}}p{{{width:.2f}\textwidth}}" for _ in range(cols)])
    out = [r"\begin{small}", rf"\begin{{longtable}}{{{spec}}}", r"\toprule"]
    for idx, row in enumerate(rows):
        padded = row + [""] * (cols - len(row))
        out.append(" & ".join(latex_inline(cell) for cell in padded) + r" \\")
        out.append(r"\midrule" if idx == 0 else "")
    out.extend([r"\bottomrule", r"\end{longtable}", r"\end{small}", ""])
    return "\n".join(line for line in out if line != "")


def build_latex(text: str, diagrams: list[dict[str, str]]) -> None:
    preamble = r"""% !TEX program = tectonic
\documentclass[UTF8,a4paper,11pt,fontset=windows]{ctexart}
\usepackage[margin=2.15cm]{geometry}
\usepackage{graphicx}
\usepackage{xcolor}
\usepackage{booktabs}
\usepackage{longtable}
\usepackage{array}
\usepackage{enumitem}
\usepackage{fancyhdr}
\usepackage{hyperref}
\usepackage{caption}
\hypersetup{colorlinks=true, linkcolor=blue, urlcolor=blue}
\setlength{\parindent}{2em}
\setlength{\parskip}{0.25em}
\setlength{\headheight}{14pt}
\setlist{nosep,leftmargin=2em}
\pagestyle{fancy}
\fancyhf{}
\lhead{AI 科研工作站框架汇报}
\rhead{\thepage}
\renewcommand{\headrulewidth}{0.4pt}
\graphicspath{{report_assets_20260615/}}
\begin{document}
\begin{titlepage}
\centering
\vspace*{3.2cm}
{\zihao{1}\bfseries 基于多 Agent 协同的 AI 科研工作站框架汇报\par}
\vspace{0.8cm}
{\Large 面向 Kaggle/HPC 机器学习科研任务的上下文分治与科研闭环系统\par}
\vspace{0.8cm}
{\large 阶段汇报文件 \quad PDF + Word\par}
\vspace{0.4cm}
{\large 2026-06-15\par}
\vfill
\noindent\colorbox{blue!7}{\parbox{0.92\textwidth}{\vspace{0.5em}\textbf{核心思想：}不是让单个最强模型硬撑完整科研流程，而是把研究、数据、模型、代码、HPC 训练、验证、提交和报告拆成可验证、可回滚、可复现的多 Agent 工作流。\vspace{0.5em}}}
\vfill
\end{titlepage}
"""
    out: list[str] = [preamble]
    diagram_idx = 0
    code_lang: str | None = None
    code_buffer: list[str] = []
    table_buffer: list[str] = []
    list_mode: str | None = None
    first_h1_seen = False

    def close_list() -> None:
        nonlocal list_mode
        if list_mode:
            out.append(rf"\end{{{list_mode}}}")
            list_mode = None

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            close_list()
            out.append(latex_table(table_buffer))
            table_buffer = []

    for line in text.splitlines():
        stripped = line.strip()
        if code_lang is not None:
            if stripped.startswith("```"):
                close_list()
                if code_lang == "mermaid":
                    if diagram_idx < len(diagrams):
                        info = diagrams[diagram_idx]
                        image_name = Path(info["path"]).name
                        out.extend(
                            [
                                r"\begin{figure}[htbp]",
                                r"\centering",
                                rf"\includegraphics[width=\linewidth]{{{latex_escape(image_name)}}}",
                                rf"\caption{{{latex_inline(latex_caption(info['caption']))}}}",
                                r"\end{figure}",
                                "",
                            ]
                        )
                        diagram_idx += 1
                else:
                    out.append(r"\begin{verbatim}")
                    out.append("\n".join(code_buffer))
                    out.append(r"\end{verbatim}")
                code_lang = None
                code_buffer = []
            else:
                code_buffer.append(line)
            continue

        if stripped.startswith("```"):
            flush_table()
            close_list()
            code_lang = stripped.strip("`").strip().lower() or "text"
            code_buffer = []
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table_buffer.append(line)
            continue
        flush_table()

        if not stripped:
            close_list()
            out.append("")
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            close_list()
            level = len(heading.group(1))
            title = heading.group(2).strip()
            if level == 1:
                if not first_h1_seen and title.startswith("基于多 Agent"):
                    first_h1_seen = True
                    continue
                out.append(rf"\section*{{{latex_inline(title)}}}")
            elif level == 2:
                out.append(rf"\section*{{{latex_inline(title)}}}")
            else:
                out.append(rf"\subsection*{{{latex_inline(title)}}}")
            continue

        if stripped.startswith(">"):
            close_list()
            out.append(r"\begin{quote}")
            out.append(r"\bfseries " + latex_inline(stripped.lstrip("> ").strip()))
            out.append(r"\end{quote}")
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet:
            if list_mode != "itemize":
                close_list()
                out.append(r"\begin{itemize}")
                list_mode = "itemize"
            out.append(r"\item " + latex_inline(bullet.group(1)))
            continue
        if numbered:
            if list_mode != "enumerate":
                close_list()
                out.append(r"\begin{enumerate}")
                list_mode = "enumerate"
            out.append(r"\item " + latex_inline(numbered.group(1)))
            continue

        close_list()
        out.append(latex_inline(stripped) + "\n")

    close_list()
    flush_table()
    out.append(r"\end{document}")
    TEX_PATH.write_text("\n".join(out), encoding="utf-8")


def main() -> int:
    if not SOURCE_MD.exists():
        raise FileNotFoundError(SOURCE_MD)
    text = SOURCE_MD.read_text(encoding="utf-8")
    diagrams = create_diagrams()
    build_docx(text, diagrams)
    build_latex(text, diagrams)
    summary = {
        "source": str(SOURCE_MD),
        "docx": str(DOCX_PATH),
        "tex": str(TEX_PATH),
        "assets": diagrams,
        "source_chars": len(text),
        "mermaid_blocks": text.count("```mermaid"),
    }
    SUMMARY_PATH.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
