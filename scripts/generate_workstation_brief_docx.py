from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_EN = ROOT / "reports" / "AI_research_workstation_architecture_brief.docx"
OUT_ZH = ROOT / "reports" / "AI科研工作站系统功能架构汇报简版.docx"


def set_font(run, size=None, bold=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = text
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def main():
    OUT_EN.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "List Bullet"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.color.rgb = RGBColor(0, 0, 0)
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.styles["Title"].font.size = Pt(18)
    doc.styles["Heading 1"].font.size = Pt(14)
    doc.styles["Heading 2"].font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AI 科研工作站系统功能架构汇报（简版）")
    set_font(run, 18, True)

    for text in [
        "汇报入口：http://127.0.0.1:8088/?page=overview",
        "负责人：景浩伟 | 日期：2026-06-28",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(text), 10.5, False)

    def h1(text):
        doc.add_paragraph(text, style="Heading 1")

    def para(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.add_run(text)

    h1("一、一句话介绍")
    para(
        "本系统是一个面向 Kaggle / MLE-Bench / HPC 机器学习任务的 AI 科研工作站。"
        "它不是单个训练脚本，而是通过多 Agent 分工、实验台账、GPU 集群调度、"
        "质量门禁、Kaggle 反馈和报告生成，形成可审计、可复现、可持续优化的科研闭环。"
    )

    h1("二、系统要解决的问题")
    add_bullets(
        doc,
        [
            "实验多以后容易丢失上下文，难以追踪模型、特征、参数和提交文件。",
            "CV 分数和 Kaggle public 分数可能不一致，需要审计和风险提示。",
            "AI 生成代码速度快，但如果没有 Gate、日志和 artifact，难以用于正式科研流程。",
            "GPU 集群使用需要规范化，避免脚本和结果散落在主目录。",
            "最终汇报需要可验证证据，而不是只给口头结果。",
        ],
    )

    h1("三、四层功能架构")
    add_table(
        doc,
        ["层级", "模块", "作用"],
        [
            ("Layer 1", "Workstation Execution 工作站执行层", "前端 UI、API、Agent 调度、GPU/HPC、Kaggle、实验台账、报告生成。"),
            ("Layer 2", "MLEvolve Search + Retrospective Memory 搜索与记忆层", "维护实验搜索图，记录成功/失败经验，支持特征工程、模型路线和融合路线持续优化。"),
            ("Layer 3", "XCIENTIST Harness 科研审计层", "为每次实验生成假设、验证合约、风险检查和 Claim Audit，防止过度宣称。"),
            ("Layer 4", "Kaggle Feedback + Island Model 反馈与并行探索层", "将 Kaggle 分数回写系统，停滞时触发特征、模型、融合等并行探索方向。"),
        ],
    )

    h1("四、多 Agent 职责分工")
    add_bullets(
        doc,
        [
            "总控 Agent：任务拆解、流程调度和结果汇总。",
            "数据审计 Agent：检查数据版本、缺失值、分布差异和泄露风险。",
            "特征工程 Agent：设计和评估新特征、变换和消融实验。",
            "模型选择 Agent：选择 CatBoost、LightGBM、XGBoost、神经网络或集成路线。",
            "Code Agent：通过 DeepSeek / Claude Code 风格生成代码草稿、diff 和 transcript。",
            "HPC/GPU Agent：通过白名单模板调度远程训练，回传日志、metrics、OOF 和 submission。",
            "验证分析 Agent：分析 CV、OOF、错误样本和稳定性。",
            "Submission Gate Agent：检查 submission schema、提交风险和人工授权。",
            "Claim Audit Agent：判断报告结论是否有充分证据支持。",
            "Report Agent：生成阶段报告、教师汇报和最终证据包。",
        ],
    )

    h1("五、当前能力与真实状态")
    add_table(
        doc,
        ["能力", "状态", "说明"],
        [
            ("工作站页面", "可访问", "http://127.0.0.1:8088/?page=overview 是系统总控台。"),
            ("DeepSeek Code Agent", "已接入", "用于代码生成、patch 草稿、会话记录和缓存优化。"),
            ("Kaggle API", "已配置", "通过 DPAPI 管理凭据；官方提交仍需要 Human Gate。"),
            ("HPC/GPU", "通道已关闭", "按规范只保留 87384 和 87318 两个允许作业，训练前必须重新 smoke。"),
            ("远端目录规范", "已收口", "所有新任务写入 /hpc2hdd/home/aimslab/research_agent_workstation。"),
            ("CatBoost", "已使用", "系统中已有 CatBoost 单模型、LGB/XGB/CatBoost 融合和依赖门禁记录。"),
        ],
    )

    h1("六、Overview 页面汇报顺序")
    add_bullets(
        doc,
        [
            "先讲 Overview：这是系统总控台，展示任务、Agent、资源、Gate 和 Evidence。",
            "再讲 Agent Runtime：每个 Agent 有明确职责，不是装饰性页面。",
            "再讲 Evidence Ledger：每个结论都必须绑定 artifact。",
            "再讲 GPU/HPC：集群调用受白名单、Gate 和专属目录约束。",
            "再讲 Integrity Gate：训练、提交和报告结论都需要质量门禁。",
            "最后讲 Report Studio：系统能够把训练结果整理成可汇报材料。",
        ],
    )

    h1("七、可以对老师这样概括")
    para(
        "老师，我现在做的是一个面向 Kaggle/MLE-Bench 的 AI 科研工作站。系统核心是四层架构："
        "底层是工作站执行和 GPU/Kaggle 调度，中间是 MLEvolve 风格的搜索和记忆层，"
        "上层是 XCIENTIST 风格的科研审计，再往上是 Kaggle 分数反馈和 Island Model 并行优化。"
        "每个实验都会产生 metrics、OOF、submission、日志和 artifact，并经过 Gate 和 Claim Audit，"
        "避免只看 public score 或过度宣称结果。现在系统已经能展示完整工作流、Agent 分工、"
        "证据台账、GPU/HPC 管理和报告生成，后续目标是把更多 Kaggle 任务和多模态任务纳入统一 benchmark，"
        "持续提升有效提交率和奖牌率。"
    )

    h1("八、汇报边界")
    add_bullets(
        doc,
        [
            "可以说：系统目标是提高 medal rate，并已有阶段性 Kaggle 自动化实验成果。",
            "不建议说：已经超过 MLEvolve 或已经完成 75 个 MLE-Bench 任务。",
            "稳妥表述：目标对齐 MLE-Bench 75 任务，后续持续追踪 valid submission rate、medal rate、reproducibility 和 auditability。",
        ],
    )

    doc.core_properties.title = "AI 科研工作站系统功能架构汇报简版"
    doc.core_properties.author = "景浩伟"
    doc.save(OUT_EN)
    doc.save(OUT_ZH)
    print(OUT_ZH)


if __name__ == "__main__":
    main()
