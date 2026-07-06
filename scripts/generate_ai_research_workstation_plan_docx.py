from __future__ import annotations

from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZipFile, ZIP_DEFLATED
import re

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports"
DOCX_PATH = OUT_DIR / "可审计自进化AI科研工作站研究计划书.docx"

BLACK = RGBColor(0, 0, 0)


def set_cell_shading(cell, fill: str = "F2F2F2") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = "000000", size: str = "4") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_font(run, name: str = "宋体", size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.color.rgb = BLACK
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def configure_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.first_line_indent = Pt(21)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for name, size, before, after in [
        ("Heading 1", 15, 12, 6),
        ("Heading 2", 13, 8, 4),
        ("Heading 3", 12, 6, 3),
    ]:
        style = styles[name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLACK
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def setup_page(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("可审计自进化AI科研工作站研究计划书")
    set_font(run, "宋体", 9, False)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("第 ")
    set_font(run, "宋体", 9, False)
    add_page_number(footer)
    run = footer.add_run(" 页")
    set_font(run, "宋体", 9, False)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("可审计自进化 AI 科研工作站研究计划书")
    set_font(run, "黑体", 20, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("面向 Kaggle / MLE-Bench 的多智能体机器学习工程系统")
    set_font(run, "黑体", 14, True)

    meta = [
        ("项目类别", "人工智能科研工具与机器学习工程平台"),
        ("申请方向", "智能体系统、AutoML、机器学习工程、科研可审计平台"),
        ("申请人", "景浩伟"),
        ("形成日期", "2026年7月3日"),
        ("文档用途", "学术申请、项目汇报与后续研究立项"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(4.0)
    table.columns[1].width = Cm(9.5)
    for key, value in meta:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
        for idx, cell in enumerate(cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
            for para in cell.paragraphs:
                para.paragraph_format.first_line_indent = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    set_font(run, "宋体", 10.5, idx == 0)
        set_cell_shading(cells[0], "F2F2F2")

    doc.add_paragraph()
    doc.add_page_break()


def h1(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 1")


def h2(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 2")


def para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_font(run, "宋体", 10.5, False)


def no_indent_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_font(run, "宋体", 10.5, bold)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(21)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run("• " + item)
        set_font(run, "宋体", 10.5, False)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    hdr_cells = table.rows[0].cells
    for idx, head in enumerate(headers):
        hdr_cells[idx].text = head
        set_cell_shading(hdr_cells[idx], "F2F2F2")
        set_cell_border(hdr_cells[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            set_cell_border(cells[idx])
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                for run in p.runs:
                    set_font(run, "宋体", 9.5, row is table.rows[0])
    doc.add_paragraph()


def build_document() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    setup_page(doc)
    add_title(doc)

    h1(doc, "摘要")
    para(
        doc,
        "本项目拟构建一个面向 Kaggle、MLE-Bench 与 HPC/GPU 机器学习任务的可审计自进化 AI 科研工作站。"
        "系统以多智能体协同、artifact-based workflow、实验台账、可复现报告和提交门禁为基础，"
        "进一步吸收 MLEvolve 的自进化搜索思想与 XCIENTIST 的科研验证框架，形成“执行层、搜索层、验证层、记忆与评测层”四层架构。"
        "研究目标不是把机器学习竞赛简化为一次性脚本训练，而是将任务解析、数据审计、代码生成、GPU/HPC 执行、CV/OOF 验证、Kaggle 提交、证据追踪和报告审计整合为长期运行的科研操作系统。"
    )
    para(
        doc,
        "项目的长期评测对象为 MLE-Bench 风格的 75 个 Kaggle 任务。系统将持续记录 valid submission rate、medal rate、top30 rate、best score trajectory、reproducibility score、auditability score 与 claim drift rate。"
        "在证据不足时，系统只允许输出 proxy evaluation 或 preliminary result，不允许虚构奖牌、排名或超过既有系统的结论。"
        "本计划书用于说明研究问题、技术路线、阶段目标、可行性基础、风险控制和预期成果。"
    )

    h1(doc, "一、研究背景与问题提出")
    para(
        doc,
        "机器学习工程任务已经从单模型训练逐渐演化为包含数据理解、特征工程、验证策略、计算资源管理、实验记录、提交审查和结果解释的复杂流程。"
        "Kaggle 与 MLE-Bench 类任务为评估 AI Agent 的机器学习工程能力提供了可量化场景，但现有自动化系统仍普遍存在三个问题："
        "第一，训练脚本与实验证据分离，导致结果难以复现；第二，搜索过程缺乏跨分支记忆，容易重复失败路线；第三，系统容易把 public leaderboard 分数误写成科研结论，形成 claim drift。"
    )
    para(
        doc,
        "本项目提出的 AI 科研工作站将上述问题转化为一个工程化研究问题：如何构建一个能够自主拆解任务、调度智能体、调用算力、生成实验、积累记忆、审计结论并对齐标准评测的科研操作系统。"
        "系统既要具备自动优化能力，又必须保留人工可审查的证据链；既要追求有效提交率和奖牌率，也要防止数据泄漏、过拟合 public leaderboard 和未证实的成果宣传。"
    )

    h1(doc, "二、国内外研究现状与不足")
    h2(doc, "2.1 MLE-Bench 的评测启示")
    para(
        doc,
        "MLE-Bench 将 75 个 Kaggle 机器学习工程竞赛整理为 AI Agent 评测基准，覆盖数据准备、建模、实验、提交和评分等真实流程。"
        "该基准的重要价值在于把“能否独立完成机器学习工程任务”转化为 valid submission、leaderboard score 与 medal-level performance 等可比较指标。"
        "然而，MLE-Bench 主要是评测框架，本项目需要进一步建设可长期运行的工作站、证据管理、前端交互和跨任务记忆系统。"
    )
    h2(doc, "2.2 MLEvolve 的自进化搜索思想")
    para(
        doc,
        "MLEvolve 强调通过 Progressive MCGS、多分支搜索、Retrospective Memory 和 Base/Stepwise/Diff 代码生成模式提升长周期机器学习工程优化能力。"
        "这些思想说明，自动化 MLE 系统不应只顺序尝试固定脚本，而应维护搜索图、允许跨分支参考，并在前期探索与后期利用之间动态切换。"
        "本项目计划将这些机制嵌入工作站 Search Controller，使每一次失败和成功都转化为下一轮可复用经验。"
    )
    h2(doc, "2.3 XCIENTIST 的科研验证框架")
    para(
        doc,
        "XCIENTIST 将文献证据、研究假设、实现计划、消融记录和修复轨迹外显为可检查 artifact，并提出 claim drift 风险：可运行产物可能已经不能支持最初声称的机制。"
        "这对本项目具有直接启示：每次实验必须先生成 validation contract，实验后必须执行 claim audit，最终报告中的每个结论都需要绑定 exp_id、metrics、artifact 和消融证据。"
    )

    h1(doc, "三、研究目标")
    bullets(
        doc,
        [
            "构建一个四层 AI 科研工作站，实现从任务接入、Agent 编排、代码生成、GPU/HPC 训练、实验审计到报告生成的闭环。",
            "实现 MLEvolve-style Search Controller，使系统具备多分支搜索、阶段切换、跨分支参考、失败归因与 retrospective memory 复用能力。",
            "实现 XCIENTIST-style Research Harness，使每个实验都有 hypothesis、implementation contract、metric、acceptance criteria、risk checklist、ablation plan 和 claim boundary。",
            "建立 MLE-Bench 75 长期评测体系，持续追踪 valid submission rate、medal/top30 rate、reproducibility、auditability 和 gap-to-target。",
            "形成可演示、可复现、可扩展的前端工作站，使用户输入 Kaggle 任务后能够在页面看到 Agent 运行、代码产物、证据、门禁、报告和资源状态。",
        ],
    )

    h1(doc, "四、研究内容与系统架构")
    para(
        doc,
        "本项目采用四层架构。第一层是 Multi-Agent Research OS，负责真实执行与证据留存；第二层是 MLEvolve-style Search Controller，负责自进化搜索与策略选择；"
        "第三层是 XCIENTIST-style Research Harness，负责验证合约与声明审计；第四层是 Memory / Benchmark / Evolution Layer，负责跨任务经验积累、长期评测和系统级进化。"
    )
    add_table(
        doc,
        ["层级", "核心职责", "主要模块或工件", "验收要点"],
        [
            [
                "Layer 1：Multi-Agent Research OS",
                "任务解析、数据审计、代码实现、HPC/GPU 执行、CV/OOF、submission 门禁、报告生成",
                "AgentOrchestrator、agent_trace、metrics、OOF、submission_audit、artifact_manifest",
                "每次 run 由工作站发起并留下完整 artifact",
            ],
            [
                "Layer 2：Search Controller",
                "多分支搜索图、Progressive MCGS、exploration/exploitation、Base/Stepwise/Diff 代码模式",
                "search_graph、mlevolve_controller、mcgs_selector、strategy_selector",
                "下一轮实验由结构化 decision 决定，失败不覆盖 best-so-far",
            ],
            [
                "Layer 3：Research Harness",
                "hypothesis、contract、metric、ablation、risk check、claim boundary、claim drift audit",
                "validation_contract、claim_audit、rank_promotion_gate、benchmark_claim_gate",
                "报告结论必须绑定证据，证据不足时只能输出 weak/unsupported",
            ],
            [
                "Layer 4：Memory / Benchmark",
                "跨任务记忆、失败归因、可复用策略、MLE-Bench 75 结果统计、gap report",
                "retrospective_memory、benchmark_manager、task_benchmark_state、leaderboard report",
                "不只记录成功任务，失败任务也必须纳入统计",
            ],
        ],
    )

    h2(doc, "4.1 前端工作站与人机协同")
    para(
        doc,
        "前端工作站定位为科研 OS 操作台，而非静态展示页。当前页面体系包括 Research Overview、AI Control、Experiments、Evolution Engine、Data & Kaggle、Report Studio、Code Agent IDE、GPU/HPC、Evidence Ledger、Integrity Gates、Literature、Task Queue、Agent Runtime、Workflow Graph 与 Settings。"
        "这些页面需要对应后端 API 与 artifact 目录，使用户能够看到训练任务何时发起、由哪些 Agent 执行、产生了哪些代码与数据、是否通过门禁、是否具备官方提交资格。"
    )
    h2(doc, "4.2 自进化实验流程")
    para(
        doc,
        "系统的基本闭环为：用户选择任务或导入 Kaggle 配置；工作站生成 task spec 和 data audit；Search Controller 选择 baseline 或分支策略；Code Agent 生成代码并进入质量门禁；"
        "GPU/HPC Runner 执行训练；Harness 读取 metrics、OOF 与 submission；Gate 判断 promote、hold 或 blocked；Memory 记录可复用策略与失败模式；Report Studio 生成可审计报告。"
    )

    h1(doc, "五、关键技术路线")
    add_table(
        doc,
        ["技术环节", "实现方法", "研究价值"],
        [
            ["多 Agent 上下文分治", "将任务解析、数据审计、模型设计、代码实现、训练执行、验证审查、报告写作拆分为独立角色", "降低长任务上下文混乱，提高可追踪性"],
            ["Artifact-based workflow", "每个阶段输出结构化 JSON/CSV/MD 工件，并通过 manifest 记录哈希和路径", "保证复现、回退和审计"],
            ["Progressive MCGS 搜索", "把 baseline、模型族、特征、调参、融合、消融组织为 search graph", "支持长周期自进化提分"],
            ["Retrospective Memory", "记录 what_worked、what_failed、metric_delta、failure_pattern 和 reusable_strategy", "避免重复失败，促进跨任务迁移"],
            ["Validation Contract", "实验前固化 hypothesis、implementation requirement、metric、baseline 与 acceptance criteria", "把优化尝试转化为可验证科研行为"],
            ["Claim Audit", "检查报告结论与 metrics、artifact、ablation 是否一致，识别 semantic/experimental/mechanistic drift", "防止过度宣传和 leaderboard overclaim"],
            ["Benchmark Manager", "按照 task schema 和 result schema 管理 75 任务评测，输出 gap report", "与 MLE-Bench / MLEvolve 指标对齐"],
        ],
    )

    h1(doc, "六、可行性基础与阶段性进展")
    para(
        doc,
        "项目已经具备较完整的工程基础：本地仓库包含 `src/research_agent_workstation` 执行层、`src/research_os` 进化层、`web/research-agent-workstation` 前端工作站、`configs/schemas` 数据结构、`prompts/agents` 智能体提示模板、`benchmark` 任务注册、`reports` 汇报与审计材料。"
        "前端工作站已围绕任务、实验、代码、GPU/HPC、证据、门禁、文献检索与报告生成建立页面结构，后续重点是继续把每个按钮和页面绑定到真实 API 与 artifact。"
    )
    para(
        doc,
        "阶段性运行结果显示，系统已在若干 tabular 任务上完成工作站闭环验证。2026年7月2日的 MLE-Bench 工作站运行总结中，系统以 fast 模式对 tabular-playground-series-may-2022、new-york-city-taxi-fare-prediction、tabular-playground-series-dec-2021 和 leaf-classification 进行测试，其中 3 个 tabular 任务生成了完整工件，1 个图像任务被正确标记为需要专用图像管道。"
        "这些结果可以作为系统能力验证，但由于部分结果仍属于采样 proxy 与本地 CV，不作为最终奖牌率或超过 MLEvolve 的结论。"
    )
    para(
        doc,
        "项目运行日志还显示，系统已接入 GPU/HPC、Kaggle 门禁、DeepSeek/Code Agent、报告生成与部分官方提交流程。后续研究将把这些能力统一纳入四层架构的 evidence gate，确保所有分数、排名和奖牌声明都有官方 response 或可复现 benchmark result 支持。"
    )

    h1(doc, "七、研究计划与进度安排")
    add_table(
        doc,
        ["阶段", "周期", "主要任务", "阶段成果"],
        [
            ["Phase 1：系统稳定与接口打通", "第1-2个月", "完善前端 API、任务发起、Agent trace、导出报告、GPU/HPC 状态与本地/远端算力选择", "可演示的端到端工作站闭环"],
            ["Phase 2：单任务全量闭环", "第3个月", "选择 1-3 个 tabular 任务进行全量训练、门禁、人工审批与官方提交", "可复现 submission、官方 response 和任务报告"],
            ["Phase 3：多任务自进化", "第4-5个月", "扩展到 10-15 个任务，完善 search graph、retrospective memory、失败归因和 best-so-far 保护", "benchmark summary 与 gap report"],
            ["Phase 4：多模态任务支持", "第6-8个月", "接入图像、文本、时序和科学数据任务管道，建立 modality-specific templates", "非 tabular 任务的 valid submission 或 failure artifact"],
            ["Phase 5：MLE-Bench 75 对齐", "第9-12个月", "覆盖 75 任务评测，统一预算、提交规则、统计口径和 claim audit", "75 任务 benchmark report 与对标分析"],
        ],
    )

    h1(doc, "八、预期创新点")
    bullets(
        doc,
        [
            "提出面向机器学习工程竞赛的“科研 OS”范式，将训练、提交、审计、报告和记忆统一为可运行工作站。",
            "把 MLEvolve 的自进化搜索思想落地到工程系统中，形成搜索图、跨分支参考、回溯记忆和多模式代码生成的闭环。",
            "将 XCIENTIST 的 validation contract 与 claim drift audit 引入 Kaggle/MLE-Bench 工作流，避免把代理指标误写成科研结论。",
            "建立前端可观测的多 Agent 工作台，让用户能在页面层面追踪代码、实验、证据、门禁、报告和资源状态。",
            "形成面向 75 任务的长期 benchmark manager，使系统表现可以按 valid submission、medal/top30、复现性和可审计性持续量化。",
        ],
    )

    h1(doc, "九、预期成果")
    bullets(
        doc,
        [
            "一个可本地部署并可接入 GPU/HPC 的 AI 科研工作站原型系统。",
            "一套可复用的多 Agent 工作流、Search Controller、Research Harness、Memory 与 Benchmark 代码模块。",
            "覆盖若干 Kaggle/MLE-Bench 任务的可复现实验台账、submission、OOF、metrics、claim audit 和报告材料。",
            "一份 MLE-Bench 75 长期评测报告，诚实呈现成功、失败、gap-to-MLEvolve 与下一轮优化计划。",
            "面向学术汇报、论文写作或软件著作权申请的系统文档、架构图、实验报告和演示页面。",
        ],
    )

    h1(doc, "十、风险分析与应对措施")
    add_table(
        doc,
        ["风险", "表现", "应对措施"],
        [
            ["数据与比赛权限风险", "Kaggle rules 未接受、403、数据下载不完整", "建立任务 readiness 检查，缺失任务标记 blocked，不虚构可运行状态"],
            ["训练资源风险", "HPC/GPU 连接变更、依赖缺失、作业超时", "使用专属目录、job manifest、smoke test 和 failure artifact；保留本地小样本验证模式"],
            ["过拟合与榜单风险", "CV-public gap 大、public score 被过度优化", "设置 submission gate、rank promotion gate 和 CV-public gap risk check"],
            ["声明漂移风险", "报告中 claim 超出实验支持范围", "所有结论经 claim audit；证据不足时输出 weak evidence 或 unsupported claim"],
            ["多模态扩展风险", "图像/文本任务不能用 tabular pipeline", "建立专用数据加载器、模型模板和 modality-specific agent prompt"],
            ["安全合规风险", "密钥泄露、token 被打印或写入报告", "凭据使用加密存储和只读 smoke test；报告与日志扫描 secret"],
        ],
    )

    h1(doc, "十一、科研伦理、数据安全与合规")
    para(
        doc,
        "本项目坚持可审计、可复现和不过度声明原则。所有官方 Kaggle 提交必须经过人工 approval gate；没有官方 response 或可对齐 benchmark result 时，不声明官方排名、奖牌或已经超过外部系统。"
        "系统不得泄露测试标签，不得手动篡改 leaderboard 结果，不得只报告成功任务，失败任务也必须进入台账。"
    )
    para(
        doc,
        "数据和凭据管理方面，训练脚本、下载数据和中间产物应放置在指定项目目录或远端专属目录；Kaggle token、SSH 密码、API key 不进入报告、代码仓库或前端展示。"
        "所有导出的报告只包含任务状态、指标、artifact 路径和审计结论，不包含明文密钥。"
    )

    h1(doc, "十二、结论")
    para(
        doc,
        "本研究计划面向一个明确目标：构建可长期运行、可复现、可审查并能持续进化的 AI 科研工作站。"
        "与传统 AutoML 或单次训练脚本不同，本项目把 Agent 编排、搜索优化、科研验证、记忆沉淀和 benchmark 对齐统一到一个系统中。"
        "后续研究将以 MLE-Bench 75 为长期评测基准，在诚实记录失败和限制的前提下逐步提升 valid submission rate、top30 rate 和 medal rate，最终形成可用于科研申请、课堂展示、系统论文和真实机器学习工程任务的完整平台。"
    )

    h1(doc, "参考文献与资料")
    references = [
        "[1] Chan, J. S., Chowdhury, N., Jaffe, O., et al. MLE-bench: Evaluating Machine Learning Agents on Machine Learning Engineering. arXiv:2410.07095, 2024. https://arxiv.org/abs/2410.07095",
        "[2] OpenAI. MLE-bench GitHub Repository. https://github.com/openai/mle-bench",
        "[3] Du, S., Yan, X., Shi, J., et al. MLEvolve: A Self-Evolving Framework for Automated Machine Learning Algorithm Discovery. arXiv:2606.06473, 2026. https://arxiv.org/abs/2606.06473",
        "[4] InternScience. MLEvolve GitHub Repository. https://github.com/InternScience/MLEvolve",
        "[5] Wang, Z., Li, H., Yang, Z., et al. Externalizing Research Synthesis and Validation in AI Scientists through a Research Harness. arXiv:2606.18874, 2026. https://arxiv.org/abs/2606.18874",
        "[6] 本项目本地文档：docs/THREE_LAYER_RESEARCH_OS_ARCHITECTURE.md；docs/MLE_BENCH_75_EVALUATION_PLAN.md；reports/MLEBENCH_WORKSTATION_RUN_SUMMARY_20260702.md。",
    ]
    for ref in references:
        no_indent_para(doc, ref)

    doc.save(DOCX_PATH)
    force_black_font_colors(DOCX_PATH)


def force_black_font_colors(path: Path) -> None:
    """Normalize all OOXML font color declarations to black."""
    color_pattern = re.compile(rb'(w:color\b[^>]*?\bw:val=")([^"]+)(")')
    with NamedTemporaryFile(delete=False, suffix=".docx", dir=path.parent) as tmp:
        tmp_path = Path(tmp.name)
    try:
        with ZipFile(path, "r") as zin, ZipFile(tmp_path, "w", ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                    data = color_pattern.sub(rb"\g<1>000000\g<3>", data)
                zout.writestr(item, data)
        tmp_path.replace(path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink()


if __name__ == "__main__":
    build_document()
    print(DOCX_PATH)
